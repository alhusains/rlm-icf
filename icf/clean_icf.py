"""
Draft ICF document generator (UHN publication-quality layout).

Produces the draft ICF that the study team reviews and edits before CAPCR
submission:

  - UHN logo in the header (top-left)
  - Bordered footer with automatic "Page X of Y" page numbering
  - All-caps underlined section headings (matching the approved-ICF style)
  - Justified body text in Arial 11 pt
  - Standard UHN signature pages appended verbatim (only the TITLE line changes)

Each section carries a small grey italic status/confidence annotation below its
heading. Body text is colour-coded by provenance so the study team can see at a
glance what they must not touch versus what needs review:

  - Required template wording (verbatim, must not be changed): default black.
  - Suggested template wording that was kept as-is: blue (#0563C1), matching the
    template's own suggested-text colour.
  - Everything else — content the AI wrote, drafted, or filled in itself,
    including ``[PLEASE COMPLETE]`` markers — is highlighted yellow.

No confidence colour-coding, evidence quotes, or review flags appear here —
those live in the separate marked-up ICF (see assemble.py).
"""

from __future__ import annotations

import html as html_mod
import os
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from icf.remediate_prompts import extract_locked_phrases
from icf.runtime_injections import (
    SIGNATURE_CONSENT_SECTION_ID,
    resolve_signature_consent_bullets,
    runtime_locked_phrases,
)
from icf.types import ExtractionResult, TemplateVariable

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_FONT = "Arial"
_BODY_PT = 11
_ANNOTATION_PT = 11
_SMALL_PT = 9

# RLM emits MARKER_PLEASE_COMPLETE inline; section-level labels for empty/skipped
# sections are chosen at Word render time.
MARKER_PLEASE_COMPLETE = "[PLEASE COMPLETE]"
MARKER_REQUIRED_SUGGESTED = (
    "[Please complete using the below suggested text. This is a required section.]"
)
MARKER_OPTIONAL_SUGGESTED = (
    "[Please complete using the below suggested text, if relevant to this study.]"
)
MARKER_ADD_OTHER_ORGS = (
    "[Add any other organizations with direct access to participant records, if applicable]"
)

_AI_DISCLAIMER = (
    "Parts of the initial draft of this consent form were created with help from "
    "an artificial intelligence tool developed at University Health Network to "
    "support consent form preparation. The final approved version was reviewed "
    "by the research team."
)

_CONTENT_STATUSES = {"FOUND", "PARTIAL", "STANDARD_TEXT"}
# Top-level and nested list markers used in ICF drafts (UHN templates use "o" for sub-bullets).
_BULLET_LINE_RE = re.compile(
    r"^(?P<indent>[ \t]*)"
    r"(?P<marker>[•\-–\*·]|[oO](?=\s))"
    r"\s+"
    r"(?P<content>.*)$"
)
_BULLET_GLYPHS = ("\u2022", "\u25E6", "\u25AA")  # • ◦ ▪
# Inline "• text" bullets (glyph is part of the paragraph). Do NOT use a hanging
# first-line indent — that pushes wrapped lines further right than the first
# line's text, which is the misalignment users see on long bullets.
_BULLET_BASE_INDENT_CM = 0.63
_BULLET_NEST_STEP_CM = 0.63
_INLINE_MARKER_RE = re.compile(
    rf"({re.escape(MARKER_PLEASE_COMPLETE)}|{re.escape(MARKER_ADD_OTHER_ORGS)})"
)

_TESTS_PROCEDURES_SECTION_ID = "13.6"
_TESTS_PROCEDURES_NOT_FOUND_SUGGESTED = (
    "This section uses a table layout to show the schedule/frequency of "
    "study-related activities. Please refer to template for the exact layout "
    "of the table."
)

# Study-team draft notes shown in grey italic under the status line (standard
# template procedures sub-sections only — these IDs do not exist in minimal risk).
_DRAFT_INSTRUCTION_BY_SECTION: dict[str, str] = {
    "12.1": (
        "List (in bullet format) any standard procedures (e.g. MRI, blood draws, "
        "etc.) and tests that are outside of standard of care. Include a lay "
        "explanation, if not obvious, of what each test involves."
    ),
    "12.2": (
        "If there are experimental procedures or medical tests, include this "
        "section.  Any standard procedures (e.g., MRI, blood draw, etc.) that "
        "are outside of standard of care should be included in the "
        "'non-experimental procedures' section – this section is for procedures "
        "that are experimental (e.g., being tested as part of the research)."
    ),
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_draft_docx(
    extractions: list[ExtractionResult],
    variables: list[TemplateVariable],
    output_path: str,
    logo_path: str | None = None,
    us_funded: bool = False,
    sdm: bool = False,
) -> str:
    """Generate the draft ICF for study-team review (UHN publication layout).

    Shares the full visual style of an approved ICF (Arial 11 pt, UHN header/
    footer, all-caps underlined headings, subsection grouping, signature pages):

    - No confidence-based colour coding — all text is black.
    - A small italic grey "Status: X | Confidence: Y" annotation is written
      directly below each section / sub-section heading.
    - Sections that could not be extracted show ``[PLEASE COMPLETE]`` (or a
      suggested-text variant) highlighted in yellow, followed by suggested text
      in grey italic.
    - No evidence quotes, no review flags, no review appendix.

    Args:
        extractions: All extraction results produced by the pipeline.
        variables:   All template variables in document order.
        output_path: Destination file path for the DOCX.
        logo_path:   Optional path to the UHN logo image (placed top-left in header).

    Returns:
        The resolved *output_path*.
    """
    doc = Document()
    _configure_page(doc)
    _set_document_font(doc)
    _build_header(doc, logo_path)
    _build_footer(doc)

    ext_map: dict[str, ExtractionResult] = {e.section_id: e for e in extractions}

    _write_validation_intro_page(doc)
    if us_funded:
        _write_us_summary_sections_validation(doc, variables, ext_map)
    _write_cover_page(doc, variables, ext_map)
    _write_validation_main_body(doc, variables, ext_map)
    _write_signature_pages(doc, _get_study_title(ext_map), sdm=sdm, ext_map=ext_map)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Page layout
# ---------------------------------------------------------------------------


def _configure_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Cm(1.25)
    sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.header_distance = Cm(1.27)
    sec.footer_distance = Cm(1.27)
    sec.different_first_page_header_footer = False


def _set_document_font(doc: Document) -> None:
    """Apply Arial 11 pt as the document-wide default."""
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)
    # Also patch the low-level rFonts element so the font propagates everywhere.
    rPr = normal._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _FONT)
    rFonts.set(qn("w:hAnsi"), _FONT)
    rFonts.set(qn("w:cs"), _FONT)
    rPr.insert(0, rFonts)


# ---------------------------------------------------------------------------
# Header: UHN logo
# ---------------------------------------------------------------------------


def _apply_bullet_paragraph_format(p, *, level: int = 0) -> None:
    """Indent a bullet paragraph without hanging wrap (see _BULLET_* constants)."""
    left_cm = _BULLET_BASE_INDENT_CM + (level * _BULLET_NEST_STEP_CM)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(left_cm)
    p.paragraph_format.first_line_indent = Cm(0)


def _build_header(doc: Document, logo_path: str | None) -> None:
    header = doc.sections[0].header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if logo_path and os.path.isfile(logo_path):
        run = para.add_run()
        # Width ~1.8" keeps the logo compact in the top-left corner.
        run.add_picture(logo_path, width=Inches(1.8))


# ---------------------------------------------------------------------------
# Footer: bordered paragraph with "Page X of Y"
# ---------------------------------------------------------------------------


def _build_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)

    # Add a top border to the paragraph (single rule, matching the approved ICF).
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "auto")
    pBdr.append(top)
    pPr.append(pBdr)

    # "Version date of this form: _______________ Page X of Y"
    _add_footer_run(para, "Version date of this form: _______________    Page ")
    _add_page_field(para, "PAGE")
    _add_footer_run(para, " of ")
    _add_page_field(para, "NUMPAGES")


def _add_footer_run(para, text: str) -> None:
    r = para.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(_SMALL_PT)


def _add_page_field(para, field_type: str) -> None:
    """Append a Word PAGE or NUMPAGES auto-field to *para*."""
    p_elem = para._element
    sz_val = str(int(_SMALL_PT * 2))  # half-points

    def _mk_run() -> OxmlElement:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), _FONT)
        rFonts.set(qn("w:hAnsi"), _FONT)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), sz_val)
        rPr.append(sz)
        r.append(rPr)
        return r

    # begin
    r1 = _mk_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1.append(fc1)
    p_elem.append(r1)

    # instrText
    r2 = _mk_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {field_type} "
    r2.append(it)
    p_elem.append(r2)

    # separate
    r3 = _mk_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "separate")
    r3.append(fc3)
    p_elem.append(r3)

    # cached value placeholder
    r4 = _mk_run()
    r4.find(qn("w:rPr")).append(OxmlElement("w:noProof"))
    t = OxmlElement("w:t")
    t.text = "1"
    r4.append(t)
    p_elem.append(r4)

    # end
    r5 = _mk_run()
    fc5 = OxmlElement("w:fldChar")
    fc5.set(qn("w:fldCharType"), "end")
    r5.append(fc5)
    p_elem.append(r5)


# ---------------------------------------------------------------------------
# Introductory notes page (inserted before the cover page)
# ---------------------------------------------------------------------------

_INTRO_GREY = RGBColor(0x55, 0x55, 0x55)
_INTRO_ITALIC_PT = 10.5
_INTRO_GREY_ITALIC_PT = 10
_UHN_TEMPLATES_URL = "https://intranet.uhnresearch.ca/service/documents-and-forms"


def _write_validation_intro_page(doc: Document) -> None:
    """Introductory cover page explaining how to use the AI-generated draft."""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AI-Generated ICF Draft")
    r.font.name = _FONT
    r.font.size = Pt(14)
    r.bold = True
    r.underline = True

    _add_blank(doc)

    _intro_section_label(doc, "How to Use This Document")

    _intro_body(
        doc,
        "This document was generated using AI and is intended as a starting point for "
        "your initial draft. The AI is grounded in the protocol as the primary source "
        "of truth; however, it may still produce inaccuracies or take details out of "
        "context.",
    )
    _intro_body(
        doc,
        "Before submitting, you are responsible for reviewing and updating this "
        "document to ensure it meets REB requirements.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    _intro_run(
        p,
        "Please use the UHN consent form templates as a reference for formatting and "
        "required content: ",
    )
    _add_intro_hyperlink(p, _UHN_TEMPLATES_URL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    _intro_run(
        p,
        "As this is a beta version, we highly suggest cross-referencing against the "
        "template before submitting the consent form in CAPCR.",
        italic=True,
        size=_INTRO_ITALIC_PT,
    )

    _add_blank(doc)
    _intro_section_label(doc, "What you need to do:")

    for item in [
        "Review all content and correct any errors or inconsistencies.",
        "Ensure information accurately reflects the protocol.",
        "Confirm formatting aligns with submission standards.",
    ]:
        p = doc.add_paragraph()
        _apply_bullet_paragraph_format(p)
        r = p.add_run("\u2022 " + item)
        r.font.name = _FONT
        r.font.size = Pt(_BODY_PT)

    p = doc.add_paragraph()
    _apply_bullet_paragraph_format(p)
    _intro_run(p, "\u2022 Replace bold placeholders like ", bold=True)
    _intro_run(p, "[PLEASE COMPLETE]", bold=True, highlight=True)
    _intro_run(
        p,
        " with your own content, review the rest of the yellow-highlighted text "
        "and clear its highlighting once approved, and remove any instructional "
        "text (including this cover page and ",
        bold=True,
    )
    _intro_run(
        p,
        "grey italic text",
        bold=True,
        italic=True,
        grey=True,
        size=_INTRO_GREY_ITALIC_PT,
    )
    _intro_run(p, ") before submitting.", bold=True)

    _add_blank(doc)
    _intro_section_label(doc, "Section Status")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    _intro_run(
        p,
        "Each section is labelled to show how much of the section the AI could "
        "pre-populate ",
    )
    _intro_run(p, "based on information in the protocol", italic=True)
    _intro_run(p, ":")

    status_entries: list[tuple[str, str | None]] = [
        (
            "FOUND",
            "The section was fully populated. Review for accuracy and plain-language quality.",
        ),
        ("PARTIAL", None),
        (
            "NOT_FOUND",
            "No relevant information was located in the protocol. You must complete this section, if applicable.",
        ),
        (
            "SKIPPED",
            "This information is not consistently available within a protocol, so the AI is "
            "designed to skip it. It must be filled in by you, if needed.",
        ),
    ]
    for label, desc in status_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        rl = p.add_run(label + ":  ")
        rl.font.name = _FONT
        rl.font.size = Pt(_BODY_PT)
        rl.bold = True
        if desc is not None:
            rd = p.add_run(desc)
            rd.font.name = _FONT
            rd.font.size = Pt(_BODY_PT)
        else:
            _intro_run(
                p,
                "Some information is missing. Look for [to be completed] for fields "
                "requiring your updates. The ",
            )
            _intro_run(
                p,
                "grey italic note",
                italic=True,
                grey=True,
                size=_INTRO_GREY_ITALIC_PT,
            )
            _intro_run(p, " at the end of the section explains what's missing.")

    _add_blank(doc)
    _intro_section_label(doc, "Confidence Level")

    _intro_body(
        doc,
        "It reflects how strongly the AI was able to match information to the protocol.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    _intro_run(p, "Note: This is not a statistical measure of confidence.", italic=True)

    for label, desc in [
        ("HIGH", "Information clearly matches the protocol."),
        (
            "MEDIUM",
            "Information found, but required interpretation or was partially ambiguous. "
            "Closer review is recommended.",
        ),
        ("LOW", "Weak or indirect match. The section should be carefully verified."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        rl = p.add_run(label + ":  ")
        rl.font.name = _FONT
        rl.font.size = Pt(_BODY_PT)
        rl.bold = True
        rd = p.add_run(desc)
        rd.font.name = _FONT
        rd.font.size = Pt(_BODY_PT)

    _add_blank(doc)
    _intro_section_label(doc, "Text Colour Legend")

    _intro_body(
        doc,
        "Body text is coloured to show where it came from, so you know what you "
        "can trust as-is versus what still needs your review:",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    _intro_run(p, "Black text", bold=True)
    _intro_run(
        p,
        " is required wording from the UHN template, copied verbatim. Do not "
        "reword it.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    _intro_run(p, "Blue text", bold=True, blue=True)
    _intro_run(
        p,
        " is suggested template wording the AI kept as-is. You can edit it if it "
        "does not fit your study.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    _intro_run(p, "Text highlighted in yellow", highlight=True)
    _intro_run(
        p,
        " was written or filled in by the AI, including missing-information "
        "placeholders like [PLEASE COMPLETE]. Review all of it carefully for "
        "accuracy before submitting.",
    )

    _add_blank(doc)
    _intro_body(
        doc,
        "The same colours mark section headings, based on the UHN template: a "
        "black heading is a required section that must appear in every ICF, and "
        "a blue heading is an optional section to include only if relevant to "
        "this study.",
    )

    _add_blank(doc)
    _intro_section_label(doc, "Important Notes")

    for note in [
        "This is an AI-generated draft and has not been reviewed or approved by REB "
        "or any regulatory body",
        "The study team is responsible for reviewing, verifying, and approving all content",
        "The submitted protocol was the sole source of information for the generation of this draft",
    ]:
        p = doc.add_paragraph()
        _apply_bullet_paragraph_format(p)
        r = p.add_run("\u2022 " + note)
        r.font.name = _FONT
        r.font.size = Pt(_BODY_PT)

    doc.add_page_break()


def _intro_section_label(doc: Document, text: str) -> None:
    """Bold sub-heading for intro page sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)
    r.bold = True


def _intro_body(doc: Document, text: str) -> None:
    """Regular body paragraph for the intro page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)


def _intro_run(
    para,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    highlight: bool = False,
    grey: bool = False,
    blue: bool = False,
    size: float | None = None,
) -> None:
    """Styled run for mixed-format intro paragraphs."""
    r = para.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size if size is not None else _BODY_PT)
    r.bold = bold
    r.italic = italic
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if grey:
        r.font.color.rgb = _INTRO_GREY
    if blue:
        r.font.color.rgb = _SUGGESTED_TEXT_BLUE


def _add_intro_hyperlink(para, url: str) -> None:
    """Append a clickable hyperlink run to an intro paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), _FONT)
    r_fonts.set(qn("w:hAnsi"), _FONT)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(_BODY_PT * 2)))
    r_pr.append(sz)
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    r_pr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = url
    run.append(text_elem)
    hyperlink.append(run)
    para._p.append(hyperlink)


# ---------------------------------------------------------------------------
# US-funded Summary of Informed Consent Form (sections 1.x)
# ---------------------------------------------------------------------------

_US_SUMMARY_PAGE_TITLE = "Summary of Informed Consent Form"


def _write_us_summary_page_opening(doc: Document, ext_map: dict[str, ExtractionResult]) -> None:
    """Page title (centered, underlined) and study title line (left, bold label)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(_US_SUMMARY_PAGE_TITLE)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)
    r.bold = False
    r.underline = True

    study_title = _get_study_title(ext_map)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    rl = p.add_run("Study Title: ")
    rl.font.name = _FONT
    rl.font.size = Pt(_BODY_PT)
    rl.bold = True
    rv = p.add_run(study_title)
    rv.font.name = _FONT
    rv.font.size = Pt(_BODY_PT)
    rv.bold = False


def _write_us_summary_sections_validation(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
) -> None:
    """US-funded summary block (sections 1.x) for the draft ICF."""
    _write_us_summary_page_opening(doc, ext_map)
    _write_us_summary_section_blocks(doc, variables, ext_map)
    doc.add_page_break()


def _write_us_summary_section_blocks(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
) -> None:
    """Write 1.1 body only (no extra heading), then 1.2–1.7 with underlined headings.

    STUDY PURPOSE through ALTERNATIVES use the same style as main ICF section
    headings: uppercase, underlined, not bold. Study Title (cover-page style)
    is written separately in ``_write_us_summary_page_opening`` as bold only.
    """
    summary_vars = [v for v in variables if v.section_id.startswith("1.")]
    var_11 = next((v for v in summary_vars if v.section_id == "1.1"), None)
    rest_vars = [v for v in summary_vars if v.section_id != "1.1"]

    if var_11 is not None:
        _write_us_summary_1_1_validation(doc, var_11, ext_map.get("1.1"))

    last_sub_section: str | None = None
    for var in rest_vars:
        ext = ext_map.get(var.section_id)

        content = _get_section_content(ext)
        if not content and not var.required:
            continue

        if var.sub_section and var.sub_section != last_sub_section:
            # Uppercase + underline (not bold) — matches approved US-summary style.
            sub_required = _any_required(rest_vars, sub_section=var.sub_section)
            _add_heading(doc, var.sub_section, color=_heading_color(sub_required))
            last_sub_section = var.sub_section
        elif not var.sub_section:
            last_sub_section = None

        if ext is not None and ext.status in (
            "FOUND",
            "PARTIAL",
            "NOT_FOUND",
            "ERROR",
            "SKIPPED",
        ):
            _add_validation_annotation(doc, ext)
        if content:
            required_phrases, suggested_phrases = _template_phrase_sources(var, content)
            _add_content_block(doc, content, required_phrases, suggested_phrases)
            if ext is not None and ext.status == "PARTIAL" and ext.notes:
                _add_partial_notes(doc, ext.notes)
        else:
            _add_validation_placeholder(doc, ext, var, optional=not var.required)


def _write_us_summary_1_1_validation(
    doc: Document,
    var: TemplateVariable,
    ext: ExtractionResult | None,
) -> None:
    """Section 1.1 for validation ICF — content only, no heading."""
    content = _get_section_content(ext)
    if not content and not var.required:
        return

    if ext is not None and ext.status in (
        "FOUND",
        "PARTIAL",
        "NOT_FOUND",
        "ERROR",
        "SKIPPED",
    ):
        _add_validation_annotation(doc, ext)

    if content:
        required_phrases, suggested_phrases = _template_phrase_sources(var, content)
        _add_content_block(doc, content, required_phrases, suggested_phrases)
        if ext is not None and ext.status == "PARTIAL" and ext.notes:
            _add_partial_notes(doc, ext.notes)
    else:
        _add_validation_placeholder(doc, ext, var, optional=not var.required)


# ---------------------------------------------------------------------------
# Cover page (sections 2.x)
# ---------------------------------------------------------------------------


def _write_cover_page(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
) -> None:
    # Main ICF title — centred, bold
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Informed Consent Form for Participation in a Research Study")
    run.font.name = _FONT
    run.font.size = Pt(_BODY_PT)
    run.bold = True

    _add_blank(doc)

    # Cover fields with extracted content, or empty fields that need study-team input.
    cover_vars = [v for v in variables if v.section_id.startswith("2.")]
    for var in cover_vars:
        ext = ext_map.get(var.section_id)
        content = _get_section_content(ext)
        if not content and not _should_render_empty_section(var, ext):
            continue

        label = var.sub_section or ""
        if content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            required_phrases, suggested_phrases = _template_phrase_sources(var, content)

            if label:
                # Strip any leading repetition of the label from the extracted content
                # (e.g. "Study Title:" label + "Study Title: XYZ" content → "XYZ").
                content = _strip_label_prefix(content, label)
                rl = p.add_run(label + " ")
                rl.bold = True
                rl.font.name = _FONT
                rl.font.size = Pt(_BODY_PT)
                _add_categorized_runs(p, content, required_phrases, suggested_phrases, Pt(_BODY_PT))
            else:
                _add_categorized_runs(p, content, required_phrases, suggested_phrases, Pt(_BODY_PT))
        else:
            if label:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                rl = p.add_run(label)
                rl.bold = True
                rl.font.name = _FONT
                rl.font.size = Pt(_BODY_PT)
            _add_validation_placeholder(doc, ext, var, optional=not var.required)

        _add_blank(doc)


# ---------------------------------------------------------------------------
# Body sections (sections 3+)
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Draft body (no colour coding; grey status annotations)
# ---------------------------------------------------------------------------

_ANNOTATION_GREY = RGBColor(0x88, 0x88, 0x88)
# Matches the template's own suggested-text colour (Word "Hyperlink"-style blue).
_SUGGESTED_TEXT_BLUE = RGBColor(0x05, 0x63, 0xC1)


def _heading_color(required: bool) -> RGBColor | None:
    """Colour for a section/sub-section heading based on the registry's ``required`` flag.

    None (default black) for REQUIRED sections -- must appear in every ICF, same
    colour as required body text. Blue for OPTIONAL sections -- include only if
    relevant to this study -- matching the suggested-text colour, since "include
    if relevant" is the same kind of judgment call as "keep this suggested wording
    if it fits".
    """
    return None if required else _SUGGESTED_TEXT_BLUE


def _any_required(variables: list[TemplateVariable], **filters: str | None) -> bool:
    """True if any variable matching all *filters* (by attribute name) is required.

    Headings/sub-headings can be shared by several registry rows (e.g. one
    heading spanning sections 12, 12.1, 12.2, ...); the heading is coloured as
    required if ANY row under it is required, since the group must then appear.
    """
    return any(
        v.required and all(getattr(v, attr) == value for attr, value in filters.items())
        for v in variables
    )


# Below this length, a phrase is unreliable as a colour-coding anchor: short
# closed-choice alternatives extracted by extract_locked_phrases (e.g. "study",
# "will", "may" from a {{study/clinical trial}} or {{will/may}} placeholder)
# recur constantly as ordinary words throughout AI-generated prose elsewhere in
# the same section, and text.find() would match every one of those incidental
# occurrences, not just the placeholder's actual slot -- painting random words
# black/blue in the middle of AI-drafted sentences. extract_locked_phrases
# itself is unaffected (see remediate.py's _locked_phrases_for): it only checks
# presence, not position, so short phrases remain fully validated there.
_MIN_COLOR_PHRASE_LEN = 12


def _template_phrase_sources(
    var: TemplateVariable | None, filled_text: str
) -> tuple[list[str], list[str]]:
    """Return (required_phrases, suggested_phrases) expected verbatim in filled_text.

    Reuses extract_locked_phrases (icf/remediate_prompts.py) -- the same literal
    phrase extraction Stage 9 remediation uses to verify required/suggested
    template wording survives edits. Here it drives colour-coding instead: any
    part of the final text NOT covered by one of these phrases was written or
    filled in by the AI, not copied from the template. Phrases too short to be
    a reliable position anchor are dropped -- see _MIN_COLOR_PHRASE_LEN.

    Runtime-injected verbatim paragraphs (e.g. the SDM opening paragraph in
    section 3) are not stored in required_text/suggested_text -- they are
    mandated fixed wording added at runtime -- so they are folded into the
    required (black) set here, same as collect_section_locked_phrases does
    for review/remediation's locked-phrase protection.
    """
    if var is None:
        return [], []
    required = extract_locked_phrases(var.required_text, filled_text)
    for p in runtime_locked_phrases(var, filled_text):
        if p not in required:
            required.append(p)
    suggested = extract_locked_phrases(var.suggested_text, filled_text)
    required_phrases = [p for p in required if len(p) >= _MIN_COLOR_PHRASE_LEN]
    suggested_phrases = [p for p in suggested if len(p) >= _MIN_COLOR_PHRASE_LEN]
    return required_phrases, suggested_phrases


def _normalize_whitespace_with_offsets(text: str) -> tuple[str, list[int]]:
    """Collapse whitespace runs to single spaces, tracking original offsets.

    Returns (normalized_text, offsets) where offsets[i] is the index in the
    original *text* corresponding to normalized_text[i] (plus a trailing
    sentinel offsets[len(normalized_text)] == len(text)), so a match found in
    normalized space can be mapped back to the original text's coordinates.
    """
    chars: list[str] = []
    offsets: list[int] = []
    prev_was_space = False
    for i, ch in enumerate(text):
        if ch.isspace():
            if not prev_was_space:
                chars.append(" ")
                offsets.append(i)
            prev_was_space = True
        else:
            chars.append(ch)
            offsets.append(i)
            prev_was_space = False
    offsets.append(len(text))
    return "".join(chars), offsets


def _find_phrase_spans(text: str, phrases: list[str]) -> list[tuple[int, int]]:
    """Return non-overlapping (start, end) spans in *text* matching any *phrase*.

    Matching is whitespace-tolerant (a run of spaces/newlines in the phrase
    matches any run of whitespace in the text) -- required/suggested text is
    occasionally re-flowed with slightly different spacing by the time it
    reaches the final draft, and an exact-substring match would otherwise
    silently drop the whole phrase instead of just tolerating the drift.
    Longer matches win when candidate spans overlap (e.g. a short phrase that
    happens to be a substring of a longer one also present in the list).
    """
    normalized_text, offsets = _normalize_whitespace_with_offsets(text)

    candidates: list[tuple[int, int]] = []
    for phrase in phrases:
        p = " ".join(phrase.split())
        if not p:
            continue
        start = 0
        while True:
            idx = normalized_text.find(p, start)
            if idx == -1:
                break
            candidates.append((offsets[idx], offsets[idx + len(p)]))
            start = idx + 1

    if not candidates:
        return []

    candidates.sort(key=lambda span: (-(span[1] - span[0]), span[0]))
    occupied = bytearray(len(text))
    accepted: list[tuple[int, int]] = []
    for start, end in candidates:
        if any(occupied[start:end]):
            continue
        accepted.append((start, end))
        occupied[start:end] = bytes([1]) * (end - start)

    accepted.sort()
    return accepted


def _categorize_spans(
    text: str, required_phrases: list[str], suggested_phrases: list[str]
) -> list[tuple[int, int, str]]:
    """Return sorted, non-overlapping (start, end, category) spans over *text*.

    category is "required" or "suggested". Required always wins where the two
    would overlap -- required wording must never be recoloured as suggested.
    Gaps between returned spans are AI-generated content (handled by the caller).
    """
    required_spans = _find_phrase_spans(text, required_phrases)
    suggested_spans = _find_phrase_spans(text, suggested_phrases)

    occupied = bytearray(len(text))
    for start, end in required_spans:
        occupied[start:end] = bytes([1]) * (end - start)

    spans = [(start, end, "required") for start, end in required_spans]
    for start, end in suggested_spans:
        if any(occupied[start:end]):
            continue
        spans.append((start, end, "suggested"))

    spans.sort()
    return spans


def _add_partial_notes(doc: Document, notes: str) -> None:
    """Grey italic status explanation rendered after section content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[PARTIAL] {notes}")
    r.font.name = _FONT
    r.font.size = Pt(_ANNOTATION_PT)
    r.italic = True
    r.font.color.rgb = _ANNOTATION_GREY


def _add_validation_annotation(doc: Document, ext: ExtractionResult) -> None:
    """Render a small italic grey status/confidence line below a heading."""
    parts = [f"Status: {ext.status}"]
    if ext.confidence and ext.confidence not in ("N/A", ""):
        parts.append(f"Confidence: {ext.confidence}")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("  |  ".join(parts))
    r.font.name = _FONT
    r.font.size = Pt(_ANNOTATION_PT)
    r.italic = True
    r.font.color.rgb = _ANNOTATION_GREY


def _add_draft_section_instruction(doc: Document, var: TemplateVariable) -> None:
    """Grey italic study-team note under the status line for selected sections."""
    text = _DRAFT_INSTRUCTION_BY_SECTION.get(var.section_id, "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(_ANNOTATION_PT)
    r.italic = True
    r.font.color.rgb = _ANNOTATION_GREY


def _add_validation_placeholder(
    doc: Document,
    ext: ExtractionResult | None,
    var: TemplateVariable | None,
    optional: bool = False,
) -> None:
    """Bold yellow placeholder label + optional suggested text in grey italic."""
    label = _resolve_section_placeholder_label(var, optional=optional)
    _add_highlighted_placeholder(doc, label)

    # Co-Investigators: surface the UHN disclaimer (instructions) before suggested text
    # so the study team sees that listing co-investigators is generally discouraged.
    if var is not None and var.section_id == "2.4" and (var.instructions or "").strip():
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_inst.paragraph_format.space_before = Pt(2)
        p_inst.paragraph_format.space_after = Pt(2)
        p_inst.paragraph_format.left_indent = Cm(0.5)
        r_inst = p_inst.add_run(var.instructions.strip())
        r_inst.font.name = _FONT
        r_inst.font.size = Pt(_ANNOTATION_PT)
        r_inst.italic = True
        r_inst.font.color.rgb = _ANNOTATION_GREY

    suggested = _section_suggested_text(var, ext)
    if suggested:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.left_indent = Cm(0.5)
        r2 = p2.add_run("Suggested text: " + suggested)
        r2.font.name = _FONT
        r2.font.size = Pt(_ANNOTATION_PT)
        r2.italic = True
        r2.font.color.rgb = _ANNOTATION_GREY


def _write_validation_main_body(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
) -> None:
    body_vars = [
        v
        for v in variables
        if not v.section_id.startswith("2.")
        and not v.section_id.startswith("1.")
        and v.section_id != SIGNATURE_CONSENT_SECTION_ID
    ]
    _write_validation_body(doc, body_vars, ext_map)


def _write_validation_body(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
) -> None:
    """Write body sections for the validation-phase ICF."""
    last_heading: str | None = None
    last_sub_section: str | None = None

    for var in variables:
        ext = ext_map.get(var.section_id)

        content = _get_section_content(ext)

        # Keep optional sections that have a meaningful extraction status so the
        # EC reviewer can see them and decide whether to fill them manually.
        # Only silently drop optional sections that were never attempted (ext is
        # None) or have an uninformative status (e.g. STANDARD_TEXT already
        # handled above).
        if not content and not var.required:
            no_useful_status = ext is None or ext.status not in (
                "NOT_FOUND",
                "SKIPPED",
                "ERROR",
            )
            if no_useful_status:
                continue

        # ---- Heading -----------------------------------------------------------
        if var.heading != last_heading:
            if last_heading is not None:
                _add_blank(doc)
            heading_required = _any_required(variables, heading=var.heading)
            _add_heading(doc, var.heading, color=_heading_color(heading_required))
            last_heading = var.heading
            last_sub_section = None

        # ---- Sub-section -------------------------------------------------------
        if var.sub_section and var.sub_section != last_sub_section:
            sub_required = _any_required(
                variables, heading=var.heading, sub_section=var.sub_section
            )
            _add_subsection_heading(doc, var.sub_section, color=_heading_color(sub_required))
            last_sub_section = var.sub_section
        elif not var.sub_section:
            last_sub_section = None

        # ---- Annotation (status / confidence) ----------------------------------
        if ext is not None and ext.status in (
            "FOUND",
            "PARTIAL",
            "NOT_FOUND",
            "ERROR",
            "SKIPPED",
        ):
            _add_validation_annotation(doc, ext)

        # ---- Study-team draft instruction (selected sections only) -------------
        _add_draft_section_instruction(doc, var)

        # ---- Content -----------------------------------------------------------
        if content:
            required_phrases, suggested_phrases = _template_phrase_sources(var, content)
            _add_content_block(doc, content, required_phrases, suggested_phrases)
            if ext is not None and ext.status == "PARTIAL" and ext.notes:
                _add_partial_notes(doc, ext.notes)
        else:
            _add_validation_placeholder(doc, ext, var, optional=not var.required)


# ---------------------------------------------------------------------------
# Signature pages (standard UHN — only TITLE line changes)
# ---------------------------------------------------------------------------


def _write_signature_pages(
    doc: Document,
    study_title: str,
    *,
    sdm: bool = False,
    ext_map: dict[str, ExtractionResult] | None = None,
) -> None:
    doc.add_page_break()

    # TITLE line: "TITLE:" (plain) + " [title]" (bold) — left-aligned, 9 pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "TITLE:", bold=False, size=9)
    _run(p, " " + study_title, bold=True, size=9)

    _add_blank(doc)
    _add_blank(doc)

    # CONSENT heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "CONSENT")

    # Consent bullets sit flush with the CONSENT heading (no list indent);
    # extra space after the bullet glyph for readability.
    consent_items = resolve_signature_consent_bullets(ext_map or {}, sdm=sdm)
    for item in consent_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        _run(p, "\u2022  " + item)

    # Blank spacers before signature blocks
    for _ in range(3):
        _add_blank(doc)

    # --- Signature block 1: Participant (and SDM when enabled) ---
    participant_sig_labels = (
        ["Signature of Participant/", "Substitute Decision-Maker"]
        if sdm
        else ["Signature of Participant"]
    )
    _sig_three_column_block(
        doc,
        [
            ("_________________________", participant_sig_labels),
            ("______________________", ["PRINTED NAME"]),
            ("______________", ["Date"]),
        ],
    )

    if sdm:
        _sig_sdm_participant_printed_name_row(doc)

    for _ in range(4):
        _add_blank(doc)

    _add_blank(doc)
    _add_blank(doc)

    # Attestation for the person conducting the consent discussion (UHN signature page).
    _body_line(
        doc,
        "I have explained to the above-named participant the nature and purpose, "
        "the potential benefits, and possible risks of participation in this study. "
        "All questions that have been raised about this study have been answered.",
    )
    _add_blank(doc)

    # --- Signature block 2: Person Conducting Consent ---
    _sig_three_column_block(
        doc,
        [
            (
                "_________________________",
                ["Signature of Person Conducting", "the Consent Discussion"],
            ),
            ("______________________", ["PRINTED NAME & ROLE"]),
            ("______________", ["Date"]),
        ],
    )

    for _ in range(2):
        _add_blank(doc)

    # --- Interpreter / Witness attestation ---
    _body_line(
        doc,
        "The following attestation must be provided if the participant is unable "
        "to read or requires an oral translation: ",
    )
    _add_blank(doc)

    # Bold instruction line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(
        p,
        "If the participant is assisted during the consent process, please check "
        "the relevant box and complete the signature space below: ",
        bold=True,
    )
    _add_blank(doc)

    # ☐ Interpreter checkbox
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r_cb = p.add_run("☐")
    r_cb.font.name = _FONT
    r_cb.font.size = Pt(14)
    _run(
        p,
        "\tThe person signing below acted as an interpreter, and attests that the "
        "study as set out in the consent form was accurately sight translated "
        "and/or interpreted, and that interpretation was provided on questions, "
        "responses and additional discussion arising from this process. ",
    )
    _add_blank(doc)

    # Interpreter signature
    _sig_three_column_block(
        doc,
        [
            ("_________________________", ["PRINT NAME", "of Interpreter"]),
            ("______________________", ["Signature"]),
            ("______________", ["Date"]),
        ],
        space_before_pt=6,
        trailing_blank=False,
    )
    _add_blank(doc)

    _body_line(doc, "______________________________________________________\t")
    _body_line(doc, "Language")

    # ☐ Witness / consent-read checkbox — same layout as interpreter checkbox
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r_cb2 = p.add_run("☐")
    r_cb2.font.name = _FONT
    r_cb2.font.size = Pt(14)
    _run(
        p,
        "\tThe consent form was read to the participant. The person signing below "
        "attests that the study as set out in this form was accurately explained "
        "to the participant, and any questions have been answered. ",
    )
    _add_blank(doc)

    # Witness signature
    _sig_three_column_block(
        doc,
        [
            ("_________________________", ["PRINT NAME", "of witness"]),
            ("______________________", ["Signature"]),
            ("______________", ["Date"]),
        ],
        space_before_pt=6,
        trailing_blank=False,
    )
    _add_blank(doc)

    _body_line(doc, "____________________________\t")
    _body_line(doc, "Relationship to Participant")

    for _ in range(2):
        _add_blank(doc)
    _add_ai_disclaimer(doc)


# ---------------------------------------------------------------------------
# Paragraph / run helpers
# ---------------------------------------------------------------------------


def _run(
    para,
    text: str,
    bold: bool = False,
    underline: bool = False,
    color: RGBColor | None = None,
    size: float | None = None,
) -> None:
    r = para.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size if size is not None else _BODY_PT)
    r.bold = bold
    r.underline = underline
    if color is not None:
        r.font.color.rgb = color


def _add_blank(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _body_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, text)


def _add_heading(doc: Document, text: str, color: RGBColor | None = None) -> None:
    """All-caps underlined heading — matching the approved ICF style.

    *color* is applied to the run text when provided (confidence colour-coding).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)
    r.bold = False
    r.underline = True
    if color is not None:
        r.font.color.rgb = color


def _add_subsection_heading(
    doc: Document, text: str, color: RGBColor | None = None
) -> None:
    """Sub-section heading.

    Plain names (e.g. 'Non-Experimental Procedures') render as bold black.
    Angle-bracket labels (e.g. '<Repositories or database studies>') are
    conditional blocks from the source template, not true headings — render
    them as bold grey italic, preserving the surrounding ``<...>``.
    """
    is_conditional_label = text.startswith("<") and text.endswith(">") and len(text) > 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)
    r.bold = True
    if is_conditional_label:
        r.italic = True
        r.font.color.rgb = _ANNOTATION_GREY
    elif color is not None:
        r.font.color.rgb = color


_TABLE_ROW_RE = re.compile(r"^\s*\|")
_TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-|:]+\|\s*$")


def _split_content(text: str) -> list[tuple[str, str]]:
    """Split text into ('text', ...) and ('table', ...) segments."""
    segments: list[tuple[str, str]] = []
    current_kind: str | None = None
    current_lines: list[str] = []

    for line in text.split("\n"):
        kind = "table" if _TABLE_ROW_RE.match(line) else "text"
        if kind != current_kind:
            if current_lines:
                segments.append((current_kind, "\n".join(current_lines)))  # type: ignore[arg-type]
            current_kind = kind
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_lines and current_kind is not None:
        segments.append((current_kind, "\n".join(current_lines)))

    return segments


def _parse_markdown_table(table_text: str) -> list[list[str]]:
    """Parse a Markdown table into rows of cell strings (first row = header)."""
    rows: list[list[str]] = []
    for line in table_text.strip().split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            continue
        if _TABLE_SEP_RE.match(line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table_block(
    doc: Document,
    rows: list[list[str]],
    required_phrases: list[str],
    suggested_phrases: list[str],
) -> None:
    """Insert a bordered Word table from parsed Markdown rows.

    The first row is rendered as a bold header with light-grey shading and no
    colour-coding (headers are structural labels, not drafted content). Data
    rows are colour-coded the same way as body text (see _add_text_runs) --
    tables in this pipeline are almost always AI-generated schedules, so most
    cell content ends up highlighted yellow, which is correct.
    """
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for ri, row_data in enumerate(rows):
        is_header = ri == 0
        for ci in range(n_cols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_table_cell_runs(
                para,
                cell_text,
                required_phrases=required_phrases,
                suggested_phrases=suggested_phrases,
                is_header=is_header,
            )
            if is_header:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tc_pr.append(shd)
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")

    _add_blank(doc)


def _render_ai_generated_segment(p, text: str, font_size: Pt) -> None:
    """Render one AI-generated (uncategorized) segment, splitting out edit markers.

    [PLEASE COMPLETE] keeps its existing bold-yellow treatment so it still
    stands out from ordinary AI-drafted prose (also yellow, but not bold).
    A segment that is pure whitespace (e.g. the single space separating two
    adjacent required/suggested phrase matches) is incidental, not content the
    AI wrote -- render it plain so a lone highlighted space doesn't appear in
    the middle of otherwise-locked text.
    """
    if not text.strip():
        r = p.add_run(text)
        r.font.name = _FONT
        r.font.size = font_size
        return

    if not (MARKER_PLEASE_COMPLETE in text or MARKER_ADD_OTHER_ORGS in text) or not (
        _INLINE_MARKER_RE.search(text)
    ):
        r = p.add_run(text)
        r.font.name = _FONT
        r.font.size = font_size
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return

    for part in _INLINE_MARKER_RE.split(text):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = _FONT
        r.font.size = font_size
        if part == MARKER_PLEASE_COMPLETE:
            r.bold = True
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif part == MARKER_ADD_OTHER_ORGS:
            r.italic = True
            r.font.color.rgb = _ANNOTATION_GREY
        else:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_categorized_runs(
    p,
    text: str,
    required_phrases: list[str],
    suggested_phrases: list[str],
    font_size: Pt,
) -> None:
    """Append runs to *p*, colour-coding *text* by provenance.

    Required template wording renders in the document's default black,
    suggested template wording that was kept renders in blue (#0563C1), and
    everything else -- content the AI wrote or filled in itself, including
    edit markers -- is highlighted yellow. Shared by body paragraphs and
    table cells (different font sizes only).
    """
    spans = _categorize_spans(text, required_phrases, suggested_phrases)
    pos = 0
    for start, end, category in spans:
        if start > pos:
            _render_ai_generated_segment(p, text[pos:start], font_size)
        run_text = text[start:end]
        r = p.add_run(run_text)
        r.font.name = _FONT
        r.font.size = font_size
        if category == "suggested":
            r.font.color.rgb = _SUGGESTED_TEXT_BLUE
        pos = end
    if pos < len(text):
        _render_ai_generated_segment(p, text[pos:], font_size)


def _add_table_cell_runs(
    p,
    text: str,
    *,
    required_phrases: list[str],
    suggested_phrases: list[str],
    is_header: bool,
) -> None:
    """Write cell text. Header cells are plain bold labels (no colour-coding);
    data cells are colour-coded by provenance like body text."""
    font_size = Pt(_BODY_PT - 1)
    if is_header:
        r = p.add_run(text)
        r.font.name = _FONT
        r.font.size = font_size
        r.bold = True
        return

    _add_categorized_runs(p, text, required_phrases, suggested_phrases, font_size)


def _add_text_runs(
    p,
    text: str,
    required_phrases: list[str],
    suggested_phrases: list[str],
) -> None:
    """Append runs to *p*, colour-coding by provenance (see _add_categorized_runs)."""
    _add_categorized_runs(p, text, required_phrases, suggested_phrases, Pt(_BODY_PT))


def _parse_bullet_line(line: str) -> tuple[int, str] | None:
    """Return (nest_level, content) for a bullet/sub-bullet line, else None.

    Nesting is inferred from leading whitespace (2 spaces or 1 tab ≈ one level).
    The UHN ``o`` / ``O`` sub-bullet marker is treated as at least level 1.
    """
    match = _BULLET_LINE_RE.match(line.rstrip())
    if not match:
        return None
    indent = match.group("indent").replace("\t", "  ")
    ws_level = len(indent) // 2
    marker = match.group("marker")
    if marker.lower() == "o":
        level = max(ws_level, 1)
    else:
        level = ws_level
    content = (match.group("content") or "").strip()
    if not content:
        return None
    return min(level, len(_BULLET_GLYPHS) - 1), content


def _add_content_block(
    doc: Document,
    text: str,
    required_phrases: list[str],
    suggested_phrases: list[str],
) -> None:
    """Write a block of extracted content, splitting on newlines.

    Segments containing Markdown tables (lines starting with |) are rendered
    as proper Word tables.  Bullet lines become indented list items, with
    nested / ``o`` sub-bullets indented one level deeper.  All other lines are
    rendered as body paragraphs, colour-coded by provenance (see
    _add_categorized_runs): required text black, suggested text blue, anything
    else -- AI-drafted content, including ``[PLEASE COMPLETE]`` -- yellow.
    """
    for kind, segment in _split_content(text):
        if kind == "table":
            rows = _parse_markdown_table(segment)
            if rows:
                _add_table_block(doc, rows, required_phrases, suggested_phrases)
                continue
        # Plain text — render line by line (preserve leading whitespace for nest level)
        for line in segment.split("\n"):
            if not line.strip():
                continue
            bullet = _parse_bullet_line(line)
            if bullet is not None:
                level, content = bullet
                glyph = _BULLET_GLYPHS[level]
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _apply_bullet_paragraph_format(p, level=level)
                # Glyph is a structural marker, not drafted content -- render it
                # plain so it isn't swept into the "gap" and highlighted yellow.
                glyph_run = p.add_run(f"{glyph} ")
                glyph_run.font.name = _FONT
                glyph_run.font.size = Pt(_BODY_PT)
                _add_text_runs(p, content, required_phrases, suggested_phrases)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                _add_text_runs(p, line.strip(), required_phrases, suggested_phrases)


def _set_table_borders_none(table) -> None:
    """Remove visible borders from a Word table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tbl_pr.append(borders)


def _set_cell_margins(cell, *, top: int = 0, bottom: int = 0) -> None:
    """Set table-cell margins in twips (dxa). Defaults in Word leave a visible row gap."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _add_sig_label_lines(cell, lines: list[str]) -> None:
    """Write left-aligned label lines directly under a signature underline."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = _FONT
        r.font.size = Pt(_BODY_PT)


def _sig_three_column_block(
    doc: Document,
    columns: list[tuple[str, list[str]]],
    *,
    space_before_pt: float = 0,
    trailing_blank: bool = True,
) -> None:
    """Three-column signature block with left-aligned labels under each underline."""
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    _set_table_borders_none(table)

    col_widths = (Inches(2.35), Inches(2.05), Inches(1.35))
    for row in table.rows:
        for ci, width in enumerate(col_widths):
            row.cells[ci].width = width

    if space_before_pt:
        table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(
            space_before_pt
        )

    for ci, (underline, labels) in enumerate(columns):
        ul_cell = table.rows[0].cells[ci]
        ul_p = ul_cell.paragraphs[0]
        ul_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ul_p.paragraph_format.space_before = Pt(0)
        ul_p.paragraph_format.space_after = Pt(0)
        _set_cell_margins(ul_cell, bottom=0)
        _run(ul_p, underline)

        label_cell = table.rows[1].cells[ci]
        _set_cell_margins(label_cell, top=0)
        _add_sig_label_lines(label_cell, labels)

    if trailing_blank:
        _add_blank(doc)


def _sig_sdm_participant_printed_name_row(doc: Document) -> None:
    """SDM-only row: prompt on the left; printed-name line under the middle underline.

    Uses the same 3-column widths as ``_sig_three_column_block`` so the underline
    and label line up with the PRINTED NAME column above. The underline cell is
    bottom-aligned so a wrapping prompt in the left cell does not leave a gap
    between the line and its label. Middle column is slightly wider so
    "PRINTED NAME of Participant" stays on one line.
    """
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    _set_table_borders_none(table)

    # Match the signature block above, with a bit more width on the middle
    # column so the participant printed-name label does not wrap.
    col_widths = (Inches(2.2), Inches(2.4), Inches(1.15))
    for row in table.rows:
        for ci, width in enumerate(col_widths):
            row.cells[ci].width = width

    prompt_cell = table.rows[0].cells[0]
    prompt_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    _set_cell_margins(prompt_cell, top=0, bottom=0)
    prompt_p = prompt_cell.paragraphs[0]
    prompt_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    prompt_p.paragraph_format.space_before = Pt(6)
    prompt_p.paragraph_format.space_after = Pt(0)
    _run(prompt_p, "If consent is provided by Substitute Decision Maker:")

    ul_cell = table.rows[0].cells[1]
    ul_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    _set_cell_margins(ul_cell, top=0, bottom=0)
    ul_p = ul_cell.paragraphs[0]
    ul_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ul_p.paragraph_format.space_before = Pt(6)
    ul_p.paragraph_format.space_after = Pt(0)
    _run(ul_p, "______________________")

    # Keep date column empty but margin-tight so row height stays minimal.
    _set_cell_margins(table.rows[0].cells[2], top=0, bottom=0)
    _set_cell_margins(table.rows[1].cells[0], top=0, bottom=0)
    _set_cell_margins(table.rows[1].cells[2], top=0, bottom=0)

    label_cell = table.rows[1].cells[1]
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(label_cell, top=0, bottom=0)
    _add_sig_label_lines(label_cell, ["PRINTED NAME of Participant"])

    _add_blank(doc)


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------


def _strip_label_prefix(content: str, label: str) -> str:
    """Remove a leading label repetition from extracted content.

    Handles cases where the LLM echoes the field name in its output, e.g.:
      label="Study Title:"  content="Study Title: Some Long Title"  → "Some Long Title"
    Matching is case-insensitive; the trailing colon and surrounding whitespace
    are consumed so the result is clean.
    """
    label_clean = label.rstrip(":").strip()
    # Try matching with optional colon after the label
    pattern = re.compile(r"^" + re.escape(label_clean) + r"\s*:?\s*", re.IGNORECASE)
    return pattern.sub("", content).strip() or content


def _plain_suggested_text(var: TemplateVariable) -> str:
    """Return suggested_text as plain text, stripping HTML tags when format is 'html'."""
    raw = html_mod.unescape(var.suggested_text)
    if var.suggested_text_format == "html":
        return re.sub(r"<[^>]+>", " ", raw).strip()
    return raw


def _resolve_section_placeholder_label(
    var: TemplateVariable | None,
    *,
    optional: bool = False,
) -> str:
    if _section_suggested_text(var):
        return MARKER_OPTIONAL_SUGGESTED if optional else MARKER_REQUIRED_SUGGESTED
    return MARKER_PLEASE_COMPLETE


def _add_highlighted_placeholder(doc: Document, label: str) -> None:
    """Bold yellow-highlighted placeholder for sections requiring study-team input."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)
    r.bold = True
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _section_suggested_text(
    var: TemplateVariable | None,
    ext: ExtractionResult | None = None,
) -> str:
    if not var:
        return ""
    if (
        var.section_id == _TESTS_PROCEDURES_SECTION_ID
        and ext is not None
        and ext.status in ("NOT_FOUND", "SKIPPED")
    ):
        return _TESTS_PROCEDURES_NOT_FOUND_SUGGESTED
    if var.suggested_text:
        return _plain_suggested_text(var).strip()
    if var.required_text:
        return var.required_text.strip()
    return ""


def _add_ai_disclaimer(doc: Document) -> None:
    """Append the AI-use statement at the end of the signature page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(_AI_DISCLAIMER)
    r.font.name = _FONT
    r.font.size = Pt(_BODY_PT)


def _should_render_empty_section(
    var: TemplateVariable, ext: ExtractionResult | None
) -> bool:
    """Whether an empty section should still appear with a placeholder."""
    if var.required:
        return True
    return ext is not None and ext.status in ("NOT_FOUND", "SKIPPED", "ERROR")


def _get_section_content(ext: ExtractionResult | None) -> str | None:
    """Return the usable text content of an extraction, or None if unavailable."""
    if ext is None or ext.status not in _CONTENT_STATUSES:
        return None
    text = (ext.filled_template or ext.answer or "").strip()
    return text if text else None


def _get_study_title(ext_map: dict[str, ExtractionResult]) -> str:
    """Return the study title from section 2.1, or a placeholder if not found."""
    ext = ext_map.get("2.1")
    if ext and ext.status in _CONTENT_STATUSES:
        raw = (ext.filled_template or ext.answer or "").strip()
        if raw:
            return _strip_label_prefix(raw, "Study Title:")
        return "[Study Title]"
    return "[Study Title]"
