"""
Human review document generator for ICF evaluation results.

Generates a colour-coded Word DOCX where reviewers can:
  - Read both the AI-generated text and the approved ground truth side by side
  - See every rubric score, grade, routing mode, evidence relevance, and reasoning
  - Write their own comments directly in the document

Usage (called from run_eval_review.py):
    generate_review_doc(
        eval_report_path="output/eval_report_combined_rlm_Prot.json",
        extraction_report_path="output/extraction_report_rlm_Prot.json",
        ground_truth_path="data/approved_icf.docx",
        registry_path="data/standard_ICF_template_breakdown.json",
        output_path="output/review_rlm_Prot.docx",
    )
"""

from __future__ import annotations

import json
import os
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm


# ------------------------------------------------------------------
# Colour palette
# ------------------------------------------------------------------

_GREEN      = RGBColor(0x2E, 0x7D, 0x32)   # Excellent
_LIGHT_GREEN= RGBColor(0x66, 0xBB, 0x6A)   # Good
_YELLOW     = RGBColor(0xF9, 0xA8, 0x25)   # Borderline
_ORANGE     = RGBColor(0xEF, 0x6C, 0x00)   # Poor
_RED        = RGBColor(0xC6, 0x28, 0x28)   # Fail
_GREY       = RGBColor(0x75, 0x75, 0x75)   # N/A / skipped
_BLACK      = RGBColor(0x21, 0x21, 0x21)
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_BLUE = RGBColor(0xE3, 0xF2, 0xFD)   # section header bg
_PALE_GREY  = RGBColor(0xF5, 0xF5, 0xF5)   # table row bg

_GRADE_COLOUR = {
    "Excellent":  _GREEN,
    "Good":       _LIGHT_GREEN,
    "Borderline": _YELLOW,
    "Poor":       _ORANGE,
    "Fail":       _RED,
    "N/A":        _GREY,
    "ERROR":      _RED,
}

_GRADE_BG = {
    "Excellent":  RGBColor(0xE8, 0xF5, 0xE9),
    "Good":       RGBColor(0xF1, 0xF8, 0xE9),
    "Borderline": RGBColor(0xFF, 0xF9, 0xC4),
    "Poor":       RGBColor(0xFF, 0xE0, 0xB2),
    "Fail":       RGBColor(0xFF, 0xEB, 0xEE),
    "N/A":        RGBColor(0xF5, 0xF5, 0xF5),
    "ERROR":      RGBColor(0xFF, 0xEB, 0xEE),
}


# ------------------------------------------------------------------
# Low-level DOCX helpers
# ------------------------------------------------------------------

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    """Fill a table cell background with a solid colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _set_cell_borders(cell, border_colour: str = "BDBDBD") -> None:
    """Add thin borders around a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_colour)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para(cell, text: str, bold: bool = False, italic: bool = False,
          colour: RGBColor | None = None, size: int = 9,
          align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Add a paragraph to a cell."""
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour


def _clear_cell(cell) -> None:
    """Remove the default empty paragraph added by python-docx."""
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def _score_bar(score: float) -> str:
    """Visual score bar, e.g. '████░░░░░░ 0.70'"""
    if score < 0:
        return "N/A"
    filled = round(score * 10)
    bar = "█" * filled + "░" * (10 - filled)
    return f"{bar}  {score:.2f}"


# ------------------------------------------------------------------
# Data loading helpers
# ------------------------------------------------------------------

def _load_extraction_text(extraction_report_path: str) -> dict[str, dict]:
    """Load AI-generated text + evidence per section from extraction report."""
    with open(extraction_report_path, encoding="utf-8") as f:
        data = json.load(f)
    result = {}
    for ext in data.get("extractions", []):
        sid = ext["section_id"]
        result[sid] = {
            "text": ext.get("filled_template") or ext.get("answer") or "",
            "evidence": ext.get("evidence", []),
            "confidence": ext.get("confidence", ""),
            "status": ext.get("status", ""),
            "notes": ext.get("notes", ""),
        }
    return result


def _load_ground_truth(ground_truth_path: str, registry_path: str) -> dict[str, str]:
    """Load ground truth text per section from approved ICF DOCX."""
    if not ground_truth_path or not os.path.exists(ground_truth_path):
        return {}
    from icf.eval_ground_truth import parse_ground_truth_docx
    from icf.registry import load_template_registry
    variables = load_template_registry(registry_path)
    return parse_ground_truth_docx(ground_truth_path, variables)


def _load_eval_report(eval_report_path: str) -> dict[str, Any]:
    """Load the eval JSON report."""
    with open(eval_report_path, encoding="utf-8") as f:
        return json.load(f)


# ------------------------------------------------------------------
# Document builder
# ------------------------------------------------------------------

def generate_review_doc(
    eval_report_path: str,
    extraction_report_path: str,
    output_path: str,
    ground_truth_path: str | None = None,
    registry_path: str = "data/standard_ICF_template_breakdown.json",
) -> None:
    """Generate a human-readable review DOCX from eval + extraction reports.

    Parameters
    ----------
    eval_report_path:        Path to eval_report_combined_*.json
    extraction_report_path:  Path to extraction_report_*.json
    output_path:             Where to save the review DOCX
    ground_truth_path:       Path to approved ICF DOCX (optional)
    registry_path:           Path to ICF template registry JSON
    """
    eval_report = _load_eval_report(eval_report_path)
    ai_texts = _load_extraction_text(extraction_report_path)
    gt_texts = _load_ground_truth(ground_truth_path, registry_path) if ground_truth_path else {}

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- Cover page ----
    title = doc.add_heading("ICF Evaluation Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Evaluation report: {os.path.basename(eval_report_path)}\n"
        f"Extraction report: {os.path.basename(extraction_report_path)}\n"
        f"Ground truth: {os.path.basename(ground_truth_path) if ground_truth_path else 'Not provided'}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Colour legend
    doc.add_heading("Score Legend", 2)
    legend_table = doc.add_table(rows=1, cols=6)
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    legend_entries = [
        ("Excellent",  "Excellent ≥0.9"),
        ("Good",       "Good ≥0.7"),
        ("Borderline", "Borderline ≥0.5"),
        ("Poor",       "Poor ≥0.25"),
        ("Fail",       "Fail <0.25"),
        ("N/A",        "N/A / Skipped"),
    ]
    for idx, (grade, label) in enumerate(legend_entries):
        cell = legend_table.rows[0].cells[idx]
        _clear_cell(cell)
        _set_cell_bg(cell, _GRADE_BG[grade])
        _set_cell_borders(cell)
        _para(cell, label, bold=True, colour=_GRADE_COLOUR[grade], size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ---- Per-backend sections ----
    for backend_name, backend_data in eval_report.items():
        doc.add_heading(f"Backend: {backend_name.upper()}", 1)

        sections_data = backend_data.get("sections", [])
        averages = backend_data.get("averages", {})

        # Summary table
        doc.add_heading("Overall Rubric Averages", 2)
        if averages:
            avg_table = doc.add_table(rows=1, cols=3)
            avg_table.style = "Table Grid"
            hdr = avg_table.rows[0].cells
            for cell, txt in zip(hdr, ["Rubric", "Score", "Grade"]):
                _clear_cell(cell)
                _set_cell_bg(cell, RGBColor(0x42, 0x42, 0x42))
                _para(cell, txt, bold=True, colour=_WHITE, size=9)

            for rubric_name, avg_data in averages.items():
                row = avg_table.add_row().cells
                grade = avg_data.get("grade", "")
                score = avg_data.get("score", 0)
                _clear_cell(row[0])
                _clear_cell(row[1])
                _clear_cell(row[2])
                _set_cell_bg(row[0], _PALE_GREY)
                _set_cell_bg(row[1], _GRADE_BG.get(grade, _PALE_GREY))
                _set_cell_bg(row[2], _GRADE_BG.get(grade, _PALE_GREY))
                _set_cell_borders(row[0])
                _set_cell_borders(row[1])
                _set_cell_borders(row[2])
                _para(row[0], rubric_name, size=9)
                _para(row[1], _score_bar(score), size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[2], grade, bold=True, size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))

        doc.add_paragraph()

        # ---- Per-section pages ----
        for sec in sections_data:
            sid = sec["section_id"]
            heading = sec.get("heading", "")
            sub = sec.get("sub_section", "")
            status = sec.get("status", "")
            scores = sec.get("scores", [])

            # Section header
            section_title = f"[{sid}] {heading}"
            if sub:
                section_title += f" > {sub}"

            h = doc.add_heading(section_title, 2)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

            # Status + confidence badge
            ai_data = ai_texts.get(sid, {})
            confidence = ai_data.get("confidence", "")
            badge = f"Extraction Status: {status}   |   Confidence: {confidence}"
            bp = doc.add_paragraph(badge)
            bp.runs[0].font.size = Pt(9)
            bp.runs[0].font.color.rgb = _GREY
            bp.runs[0].italic = True

            # ---- AI Generated vs Ground Truth table ----
            doc.add_heading("Generated vs Approved ICF", 3)
            text_table = doc.add_table(rows=1, cols=2)
            text_table.style = "Table Grid"

            # Headers
            ai_hdr = text_table.rows[0].cells[0]
            gt_hdr = text_table.rows[0].cells[1]
            _clear_cell(ai_hdr)
            _clear_cell(gt_hdr)
            _set_cell_bg(ai_hdr, RGBColor(0x1A, 0x23, 0x7E))
            _set_cell_bg(gt_hdr, RGBColor(0x1B, 0x5E, 0x20))
            _para(ai_hdr, "AI Generated", bold=True, colour=_WHITE, size=10)
            _para(gt_hdr, "Approved Ground Truth", bold=True, colour=_WHITE, size=10)

            # Content row
            content_row = text_table.add_row().cells
            ai_cell = content_row[0]
            gt_cell = content_row[1]
            _clear_cell(ai_cell)
            _clear_cell(gt_cell)
            _set_cell_bg(ai_cell, RGBColor(0xE8, 0xEA, 0xF6))
            _set_cell_bg(gt_cell, RGBColor(0xE8, 0xF5, 0xE9))

            ai_text = ai_data.get("text", "").strip() or "[No text generated]"
            gt_text = gt_texts.get(sid, "").strip() or "[Not available in ground truth]"

            _para(ai_cell, ai_text, size=9)
            _para(gt_cell, gt_text, size=9)

            # ---- Evidence quotes ----
            evidence = ai_data.get("evidence", [])
            if evidence:
                doc.add_heading("Protocol Evidence Quotes", 3)
                ev_table = doc.add_table(rows=1, cols=3)
                ev_table.style = "Table Grid"
                for cell, txt in zip(ev_table.rows[0].cells, ["#", "Quote", "Location"]):
                    _clear_cell(cell)
                    _set_cell_bg(cell, RGBColor(0x42, 0x42, 0x42))
                    _para(cell, txt, bold=True, colour=_WHITE, size=8)

                for i, ev in enumerate(evidence[:8], 1):
                    row = ev_table.add_row().cells
                    quote = ev.get("quote", "").strip()
                    page = ev.get("page", "")
                    sect = ev.get("section", "")
                    loc = f"Page {page}" + (f"\n{sect}" if sect else "")
                    _clear_cell(row[0])
                    _clear_cell(row[1])
                    _clear_cell(row[2])
                    bg = _PALE_GREY if i % 2 == 0 else _WHITE
                    for c in row:
                        _set_cell_bg(c, bg)
                        _set_cell_borders(c)
                    _para(row[0], str(i), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                    _para(row[1], f'"{quote}"', size=8, italic=True)
                    _para(row[2], loc, size=8)

            # ---- Rubric scores table ----
            doc.add_heading("Rubric Evaluation", 3)
            rubric_table = doc.add_table(rows=1, cols=5)
            rubric_table.style = "Table Grid"

            for cell, txt in zip(rubric_table.rows[0].cells,
                                  ["Rubric", "Score", "Grade", "Routing / Relevance", "Reasoning"]):
                _clear_cell(cell)
                _set_cell_bg(cell, RGBColor(0x42, 0x42, 0x42))
                _para(cell, txt, bold=True, colour=_WHITE, size=9)

            for s in scores:
                grade = s.get("grade", "N/A")
                score = s.get("score", -1.0)
                rubric_name = s.get("rubric", "")
                reason = s.get("reason", "")
                routing = s.get("routing_mode", "")
                ev_rel = s.get("evidence_relevance", "")
                sup_lv = s.get("support_level", "")
                conf = s.get("confidence", "")

                row = rubric_table.add_row().cells
                for c in row:
                    _clear_cell(c)
                    _set_cell_bg(c, _GRADE_BG.get(grade, _PALE_GREY))
                    _set_cell_borders(c)

                routing_str = routing.replace("ScoringMode.", "") if routing else ""
                meta_parts = [routing_str]
                if ev_rel:
                    meta_parts.append(f"Relevance: {ev_rel}")
                if sup_lv:
                    meta_parts.append(f"Support: {sup_lv}")
                if conf:
                    meta_parts.append(f"Conf: {conf}")
                meta_str = "\n".join(meta_parts)

                _para(row[0], rubric_name, bold=True, size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[1], _score_bar(score), size=8,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[2], grade, bold=True, size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[3], meta_str, size=8, colour=_GREY)
                _para(row[4], reason, size=8)

            # ---- Reviewer comment box ----
            doc.add_heading("Reviewer Comments", 3)
            comment_table = doc.add_table(rows=1, cols=1)
            comment_table.style = "Table Grid"
            comment_cell = comment_table.rows[0].cells[0]
            _clear_cell(comment_cell)
            _set_cell_bg(comment_cell, RGBColor(0xFF, 0xFF, 0xE0))
            _set_cell_borders(comment_cell, "9E9E9E")
            _para(comment_cell, "Reviewer notes:", bold=True, size=9, colour=_GREY)
            # Three blank lines for writing space
            for _ in range(3):
                comment_cell.add_paragraph().paragraph_format.space_after = Pt(10)

            doc.add_page_break()

        # ---- Document-level scores (if present) ----
        doc_level = backend_data.get("document_level")
        if doc_level:
            doc.add_heading("Document-Level Quality", 2)
            for dl in doc_level:
                grade = dl.get("grade", "N/A")
                score = dl.get("score", -1.0)
                rubric_name = dl.get("rubric", "")
                reason = dl.get("reason", "")
                issues = dl.get("issues", [])

                dl_table = doc.add_table(rows=1, cols=3)
                dl_table.style = "Table Grid"
                for cell, txt in zip(dl_table.rows[0].cells, ["Rubric", "Score", "Grade"]):
                    _clear_cell(cell)
                    _set_cell_bg(cell, RGBColor(0x42, 0x42, 0x42))
                    _para(cell, txt, bold=True, colour=_WHITE, size=9)

                row = dl_table.add_row().cells
                for c in row:
                    _clear_cell(c)
                    _set_cell_bg(c, _GRADE_BG.get(grade, _PALE_GREY))
                    _set_cell_borders(c)
                _para(row[0], rubric_name, bold=True, size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[1], _score_bar(score), size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))
                _para(row[2], grade, bold=True, size=9,
                      colour=_GRADE_COLOUR.get(grade, _BLACK))

                # Reason
                if reason:
                    p = doc.add_paragraph()
                    run = p.add_run(reason)
                    run.font.size = Pt(9)
                    run.italic = True

                # Issues list
                if issues:
                    doc.add_heading("Issues Found", 3)
                    for issue in issues:
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(issue)
                        run.font.size = Pt(9)

            doc.add_page_break()

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"[REVIEW] Review document saved -> {output_path}")
