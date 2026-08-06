"""
Consistency evaluation: run RLM extraction N times on the same protocol and
measure how much the outputs agree.

Two entry points:
  - ``run_consistency_eval``: single-section mode. Calls
    ``ExtractionEngine.extract_variable`` directly (not the full
    ``ICFPipeline``) so the run measures pure extraction variance, without
    the Stage 5.5/8/9 harmonize/review/remediation passes adding their own
    independent LLM calls on top.
  - ``run_full_pipeline_consistency_eval``: whole-protocol mode. Runs the full
    ``ICFPipeline`` (extract -> harmonize -> review -> remediate) N times and
    compares the FINAL, post-processed extraction for every section across
    runs -- i.e. exactly what would ship in the ICF.

Two categories of measures are produced per section:
  1. Deterministic, code-only metrics computed directly from the N results
     (status/confidence agreement, pairwise text similarity, evidence overlap).
  2. An LLM-as-judge pass that reads all N results side by side and reports
     on factual/semantic agreement, dropped or added facts, and contradictions
     -- things the code-only metrics can't see.

``generate_consistency_docx`` renders a human-readable Word report (one
section per page, an executive-summary table up front) intended for a
non-technical reviewer to read and compare visually.
"""

from __future__ import annotations

import datetime
import difflib
import json
import os
import re
import time
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from icf.extract import ExtractionEngine, parse_extraction_json
from icf.ingest import load_protocol
from icf.pipeline import ICFPipeline
from icf.registry import load_template_registry
from icf.types import ExtractionResult, PipelineResult, TemplateVariable
from rlm.clients import get_client

_JUDGE_SYSTEM_PROMPT = (
    "You are auditing an AI extraction pipeline for consistency. You will be shown "
    "the SAME section of a clinical study protocol extracted N separate times by "
    "the SAME model. Your job is to judge how consistent the N outputs are with "
    "each other -- NOT whether they are correct relative to the protocol.\n\n"
    "Judge consistency along these dimensions:\n"
    "1. Do the runs agree on the core facts (numbers, names, procedures, conclusions)?\n"
    "2. Are there facts present in some runs but completely absent in others "
    "(not just reworded -- actually missing)?\n"
    "3. Are there direct contradictions between runs (e.g. different numbers "
    "for the same quantity, opposite claims)?\n"
    "4. Do the runs agree on status (FOUND/PARTIAL/NOT_FOUND) and on which "
    "protocol evidence they cite?\n\n"
    "Wording differences (paraphrasing, synonyms, sentence order) are NOT "
    "inconsistencies -- only flag them if the underlying facts differ.\n\n"
    "Respond with ONLY a JSON object (no markdown fences, no prose) with this "
    "exact shape:\n"
    "{\n"
    '  "consistency_score": <int 0-100, 100 = perfectly consistent>,\n'
    '  "consistency_label": "HIGH" | "MEDIUM" | "LOW",\n'
    '  "status_disagreement": <bool>,\n'
    '  "facts_missing_in_some_runs": [\n'
    '    {"fact": "<short description>", "present_in_runs": [1,2], "missing_in_runs": [3]}\n'
    "  ],\n"
    '  "contradictions": [\n'
    '    {"description": "<short description>", "runs_involved": [1,3]}\n'
    "  ],\n"
    '  "summary": "<2-4 sentence plain-English verdict>"\n'
    "}"
)


@dataclass
class ConsistencyRunResult:
    """One run's extraction result plus timing."""

    run_index: int
    extraction: ExtractionResult
    elapsed_seconds: float

    def to_dict(self) -> dict:
        return {
            "run_index": self.run_index,
            "elapsed_seconds": round(self.elapsed_seconds, 1),
            "extraction": self.extraction.to_dict(),
        }


@dataclass
class QuantitativeMetrics:
    """Code-only agreement metrics computed across all N runs."""

    statuses: list[str]
    confidences: list[str]
    evidence_counts: list[int]
    status_unanimous: bool
    status_majority: str
    status_agreement_rate: float
    confidence_unanimous: bool
    confidence_majority: str
    confidence_agreement_rate: float
    exact_match_pairs: int
    total_pairs: int
    avg_char_similarity: float
    min_char_similarity: float
    avg_word_jaccard: float
    min_word_jaccard: float
    avg_evidence_jaccard: float
    min_evidence_jaccard: float

    def to_dict(self) -> dict:
        return {
            "statuses": self.statuses,
            "confidences": self.confidences,
            "evidence_counts": self.evidence_counts,
            "status_unanimous": self.status_unanimous,
            "status_majority": self.status_majority,
            "status_agreement_rate": round(self.status_agreement_rate, 2),
            "confidence_unanimous": self.confidence_unanimous,
            "confidence_majority": self.confidence_majority,
            "confidence_agreement_rate": round(self.confidence_agreement_rate, 2),
            "exact_match_pairs": self.exact_match_pairs,
            "total_pairs": self.total_pairs,
            "avg_char_similarity": round(self.avg_char_similarity, 3),
            "min_char_similarity": round(self.min_char_similarity, 3),
            "avg_word_jaccard": round(self.avg_word_jaccard, 3),
            "min_word_jaccard": round(self.min_word_jaccard, 3),
            "avg_evidence_jaccard": round(self.avg_evidence_jaccard, 3),
            "min_evidence_jaccard": round(self.min_evidence_jaccard, 3),
        }


@dataclass
class JudgeVerdict:
    """Parsed (or raw-fallback) output of the LLM-as-judge pass."""

    consistency_score: int | None
    consistency_label: str
    status_disagreement: bool
    facts_missing_in_some_runs: list[dict]
    contradictions: list[dict]
    summary: str
    raw_response: str = ""

    def to_dict(self) -> dict:
        return {
            "consistency_score": self.consistency_score,
            "consistency_label": self.consistency_label,
            "status_disagreement": self.status_disagreement,
            "facts_missing_in_some_runs": self.facts_missing_in_some_runs,
            "contradictions": self.contradictions,
            "summary": self.summary,
        }


@dataclass
class ConsistencyReport:
    section_id: str
    heading: str
    protocol_path: str
    n_runs: int
    model_name: str
    seed: int | None
    runs: list[ConsistencyRunResult] = field(default_factory=list)
    metrics: QuantitativeMetrics | None = None
    judge_verdict: JudgeVerdict | None = None

    def to_dict(self) -> dict:
        return {
            "section_id": self.section_id,
            "heading": self.heading,
            "protocol_path": self.protocol_path,
            "n_runs": self.n_runs,
            "model_name": self.model_name,
            "seed": self.seed,
            "runs": [r.to_dict() for r in self.runs],
            "metrics": self.metrics.to_dict() if self.metrics else None,
            "judge_verdict": self.judge_verdict.to_dict() if self.judge_verdict else None,
        }


def _normalize_quote(quote: str) -> str:
    return " ".join(quote.lower().split())


def _word_set(text: str) -> set[str]:
    return set(re.findall(r"[a-z0-9]+", text.lower()))


def _jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 1.0
    union = a | b
    if not union:
        return 1.0
    return len(a & b) / len(union)


def _mode(values: list[str]) -> tuple[str, int]:
    """Return (most common value, its count)."""
    best_value, best_count = "", 0
    for v in set(values):
        count = values.count(v)
        if count > best_count:
            best_value, best_count = v, count
    return best_value, best_count


def compute_quantitative_metrics(runs: list[ConsistencyRunResult]) -> QuantitativeMetrics:
    n = len(runs)
    statuses = [r.extraction.status for r in runs]
    confidences = [r.extraction.confidence for r in runs]
    evidence_counts = [len(r.extraction.evidence) for r in runs]

    status_majority, status_count = _mode(statuses)
    confidence_majority, confidence_count = _mode(confidences)

    texts = [r.extraction.filled_template or "" for r in runs]
    word_sets = [_word_set(t) for t in texts]
    evidence_sets = [
        {_normalize_quote(e.quote) for e in r.extraction.evidence} for r in runs
    ]

    char_sims, word_sims, evidence_sims = [], [], []
    exact_matches, total_pairs = 0, 0
    for i in range(n):
        for j in range(i + 1, n):
            total_pairs += 1
            if texts[i] == texts[j]:
                exact_matches += 1
            char_sims.append(difflib.SequenceMatcher(None, texts[i], texts[j]).ratio())
            word_sims.append(_jaccard(word_sets[i], word_sets[j]))
            evidence_sims.append(_jaccard(evidence_sets[i], evidence_sets[j]))

    return QuantitativeMetrics(
        statuses=statuses,
        confidences=confidences,
        evidence_counts=evidence_counts,
        status_unanimous=status_count == n,
        status_majority=status_majority,
        status_agreement_rate=status_count / n,
        confidence_unanimous=confidence_count == n,
        confidence_majority=confidence_majority,
        confidence_agreement_rate=confidence_count / n,
        exact_match_pairs=exact_matches,
        total_pairs=total_pairs,
        avg_char_similarity=sum(char_sims) / len(char_sims) if char_sims else 1.0,
        min_char_similarity=min(char_sims) if char_sims else 1.0,
        avg_word_jaccard=sum(word_sims) / len(word_sims) if word_sims else 1.0,
        min_word_jaccard=min(word_sims) if word_sims else 1.0,
        avg_evidence_jaccard=sum(evidence_sims) / len(evidence_sims) if evidence_sims else 1.0,
        min_evidence_jaccard=min(evidence_sims) if evidence_sims else 1.0,
    )


def _build_judge_prompt(variable: TemplateVariable, runs: list[ConsistencyRunResult]) -> str:
    parts = [
        f"Section: [{variable.section_id}] {variable.heading}",
        "",
    ]
    for r in runs:
        ext = r.extraction
        quotes = "\n".join(f"  - {e.quote} (p.{e.page})" for e in ext.evidence) or "  (none)"
        parts.append(
            f"=== RUN {r.run_index} ===\n"
            f"status: {ext.status}\n"
            f"confidence: {ext.confidence}\n"
            f"filled_template:\n{ext.filled_template}\n"
            f"evidence quotes:\n{quotes}\n"
            f"notes: {ext.notes or '(none)'}\n"
        )
    return "\n".join(parts)


def _judge_with_llm(
    variable: TemplateVariable,
    runs: list[ConsistencyRunResult],
    judge_backend: str,
    judge_backend_kwargs: dict,
) -> JudgeVerdict:
    client = get_client(judge_backend, judge_backend_kwargs)
    prompt = [
        {"role": "system", "content": _JUDGE_SYSTEM_PROMPT},
        {"role": "user", "content": _build_judge_prompt(variable, runs)},
    ]
    raw = client.completion(prompt)
    data = parse_extraction_json(raw)

    if not isinstance(data, dict) or "consistency_score" not in data:
        return JudgeVerdict(
            consistency_score=None,
            consistency_label="UNKNOWN",
            status_disagreement=len({r.extraction.status for r in runs}) > 1,
            facts_missing_in_some_runs=[],
            contradictions=[],
            summary="Judge response could not be parsed as JSON; see raw_response.",
            raw_response=raw,
        )

    return JudgeVerdict(
        consistency_score=data.get("consistency_score"),
        consistency_label=str(data.get("consistency_label", "UNKNOWN")),
        status_disagreement=bool(data.get("status_disagreement", False)),
        facts_missing_in_some_runs=data.get("facts_missing_in_some_runs") or [],
        contradictions=data.get("contradictions") or [],
        summary=str(data.get("summary", "")),
        raw_response=raw,
    )


def _build_section_report(
    variable: TemplateVariable,
    runs: list[ConsistencyRunResult],
    protocol_path: str,
    n_runs: int,
    model_name: str,
    seed: int | None,
    backend: str,
    judge_model: str | None,
    judge_backend: str | None,
    judge_backend_kwargs: dict | None,
) -> ConsistencyReport:
    """Compute quantitative metrics + LLM-judge verdict for one section's N runs."""
    metrics = compute_quantitative_metrics(runs)

    print(f"[CONSISTENCY] [{variable.section_id}] Running LLM-as-judge pass ...")
    judge_kwargs = dict(judge_backend_kwargs or {})
    judge_kwargs["model_name"] = judge_model or model_name
    verdict = _judge_with_llm(variable, runs, judge_backend or backend, judge_kwargs)

    return ConsistencyReport(
        section_id=variable.section_id,
        heading=variable.heading,
        protocol_path=protocol_path,
        n_runs=n_runs,
        model_name=model_name,
        seed=seed,
        runs=runs,
        metrics=metrics,
        judge_verdict=verdict,
    )


def run_consistency_eval(
    protocol_path: str,
    registry_path: str,
    section_id: str,
    n_runs: int = 3,
    model_name: str = "gpt-5.4",
    backend: str = "azure_openai",
    backend_kwargs: dict | None = None,
    max_iterations: int = 20,
    verbose: bool = False,
    judge_model: str | None = None,
    judge_backend: str | None = None,
    judge_backend_kwargs: dict | None = None,
) -> ConsistencyReport:
    """Run extraction ``n_runs`` times on one section and measure consistency."""
    backend_kwargs = dict(backend_kwargs or {})

    protocol = load_protocol(protocol_path)
    variables = load_template_registry(registry_path)
    matches = [v for v in variables if v.section_id == section_id]
    if not matches:
        raise ValueError(f"Section {section_id!r} not found in registry {registry_path!r}.")
    variable = matches[0]

    if variable.is_standard_text:
        raise ValueError(
            f"Section {section_id!r} is standard boilerplate text (no extraction happens); "
            "nothing to evaluate for consistency."
        )
    if not variable.is_in_protocol and not variable.partially_in_protocol:
        raise ValueError(
            f"Section {section_id!r} is marked 'not in protocol' (always skipped); "
            "nothing to evaluate for consistency."
        )

    engine = ExtractionEngine(
        model_name=model_name,
        backend=backend,
        backend_kwargs=backend_kwargs,
        max_iterations=max_iterations,
        verbose=verbose,
    )

    runs: list[ConsistencyRunResult] = []
    for i in range(1, n_runs + 1):
        print(f"[CONSISTENCY] Run {i}/{n_runs}: extracting section {section_id} ...")
        start = time.perf_counter()
        result = engine.extract_variable(protocol.full_text, variable)
        elapsed = time.perf_counter() - start
        print(
            f"[CONSISTENCY] Run {i}/{n_runs} done in {elapsed:.1f}s -> "
            f"{result.status} / {result.confidence} / {len(result.evidence)} evidence quote(s)"
        )
        runs.append(ConsistencyRunResult(run_index=i, extraction=result, elapsed_seconds=elapsed))

    return _build_section_report(
        variable,
        runs,
        protocol_path,
        n_runs,
        model_name,
        backend_kwargs.get("seed"),
        backend,
        judge_model,
        judge_backend,
        judge_backend_kwargs or backend_kwargs,
    )


def run_full_pipeline_consistency_eval(
    protocol_path: str,
    registry_path: str,
    n_runs: int = 3,
    model_name: str = "gpt-5.4",
    backend: str = "azure_openai",
    backend_kwargs: dict | None = None,
    max_iterations: int = 20,
    verbose: bool = False,
    section_filter: list[str] | None = None,
    judge_model: str | None = None,
    judge_backend: str | None = None,
    judge_backend_kwargs: dict | None = None,
    skip_harmonize: bool = False,
    skip_review: bool = False,
    skip_remediation: bool = False,
    remediate_high_only: bool = False,
    us_funded: bool = False,
    sdm: bool = False,
    template_docx_path: str | None = None,
    pipeline_runs_dir: str = ".consistency_pipeline_runs",
) -> list[ConsistencyReport]:
    """Run the FULL ICFPipeline ``n_runs`` times (extract -> harmonize -> review ->
    remediate) and measure per-section consistency on the FINAL, post-processed
    extractions -- i.e. the same text that would ship in the ICF.

    Returns one ConsistencyReport per extractable section that produced content
    in at least 2 of the N runs, in registry order. Sections that are standard
    boilerplate or never in the protocol (SKIPPED/STANDARD_TEXT in every run)
    are omitted -- there is nothing to compare.
    """
    backend_kwargs = dict(backend_kwargs or {})

    pipeline_results: list[PipelineResult] = []
    for i in range(1, n_runs + 1):
        print(f"\n[CONSISTENCY] {'=' * 60}")
        print(f"[CONSISTENCY] FULL PIPELINE RUN {i}/{n_runs}")
        print(f"[CONSISTENCY] {'=' * 60}")
        pipeline = ICFPipeline(
            protocol_path=protocol_path,
            template_path=registry_path,
            template_docx_path=template_docx_path,
            output_dir=os.path.join(pipeline_runs_dir, f"run{i}"),
            model_name=model_name,
            backend=backend,
            backend_kwargs=backend_kwargs,
            max_iterations=max_iterations,
            verbose=verbose,
            section_filter=section_filter,
            skip_review=skip_review,
            skip_remediation=skip_remediation,
            remediate_high_only=remediate_high_only,
            skip_harmonize=skip_harmonize,
            us_funded=us_funded,
            sdm=sdm,
        )
        pipeline_results.append(pipeline.run())

    all_variables = load_template_registry(registry_path)
    variables_by_id = {v.section_id: v for v in all_variables}
    section_ids = section_filter or [v.section_id for v in all_variables]

    reports: list[ConsistencyReport] = []
    for section_id in section_ids:
        variable = variables_by_id.get(section_id)
        if variable is None or variable.is_standard_text:
            continue

        runs: list[ConsistencyRunResult] = []
        for i, presult in enumerate(pipeline_results, start=1):
            matches = [e for e in presult.extractions if e.section_id == section_id]
            if matches:
                runs.append(ConsistencyRunResult(run_index=i, extraction=matches[0], elapsed_seconds=0.0))

        if len(runs) < 2:
            continue
        if all(r.extraction.status in ("SKIPPED", "STANDARD_TEXT") for r in runs):
            continue

        print(
            f"[CONSISTENCY] Analyzing section [{section_id}] {variable.heading} "
            f"across {len(runs)} run(s) ..."
        )
        reports.append(
            _build_section_report(
                variable,
                runs,
                protocol_path,
                n_runs,
                model_name,
                backend_kwargs.get("seed"),
                backend,
                judge_model,
                judge_backend,
                judge_backend_kwargs or backend_kwargs,
            )
        )

    return reports


def print_report(report: ConsistencyReport) -> None:
    m = report.metrics
    j = report.judge_verdict
    sep = "=" * 64
    print(f"\n{sep}")
    print(f"CONSISTENCY REPORT — Section [{report.section_id}] {report.heading}")
    print(sep)
    print(f"  Runs:                {report.n_runs}")
    print(f"  Model:                {report.model_name} (seed={report.seed})")
    if m:
        print(f"  Status per run:       {m.statuses}")
        print(
            f"  Status agreement:     {m.status_agreement_rate:.0%} "
            f"({'unanimous' if m.status_unanimous else 'majority=' + m.status_majority})"
        )
        print(f"  Confidence per run:   {m.confidences}")
        print(f"  Evidence count/run:   {m.evidence_counts}")
        print(
            f"  Exact text matches:   {m.exact_match_pairs}/{m.total_pairs} run pairs "
            "byte-identical"
        )
        print(f"  Avg char similarity:  {m.avg_char_similarity:.2f} (min {m.min_char_similarity:.2f})")
        print(f"  Avg word overlap:     {m.avg_word_jaccard:.2f} (min {m.min_word_jaccard:.2f})")
        print(
            f"  Avg evidence overlap: {m.avg_evidence_jaccard:.2f} (min {m.min_evidence_jaccard:.2f})"
        )
    if j:
        print(f"\n  LLM judge score:      {j.consistency_score}/100 ({j.consistency_label})")
        print(f"  Status disagreement:  {j.status_disagreement}")
        if j.facts_missing_in_some_runs:
            print("  Facts missing in some runs:")
            for f in j.facts_missing_in_some_runs:
                print(
                    f"    - {f.get('fact')} "
                    f"(present in {f.get('present_in_runs')}, missing in {f.get('missing_in_runs')})"
                )
        if j.contradictions:
            print("  Contradictions:")
            for c in j.contradictions:
                print(f"    - {c.get('description')} (runs {c.get('runs_involved')})")
        print(f"  Judge summary:        {j.summary}")
    print(sep)


def save_report(report: ConsistencyReport, output_path: str) -> None:
    with open(output_path, "w") as f:
        json.dump(report.to_dict(), f, indent=2)
    print(f"[CONSISTENCY] Report saved -> {output_path}")


def print_multi_report_summary(reports: list[ConsistencyReport]) -> None:
    """One-line-per-section overview, meant for the run log (background-safe)."""
    sep = "=" * 72
    print(f"\n{sep}")
    print(f"CONSISTENCY SUMMARY — {len(reports)} section(s) analyzed")
    print(sep)
    for r in reports:
        j = r.judge_verdict
        score = j.consistency_score if j else None
        label = j.consistency_label if j else "N/A"
        flag = " [REVIEW]" if _needs_attention(r) else ""
        print(f"  [{r.section_id:<6}] {r.heading[:45]:<45} score={score}/100 ({label}){flag}")
    print(sep)


def save_reports_json(
    reports: list[ConsistencyReport],
    output_path: str,
    protocol_path: str,
    n_runs: int,
    model_name: str,
) -> None:
    data = {
        "protocol_path": protocol_path,
        "n_runs": n_runs,
        "model_name": model_name,
        "generated_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "sections": [r.to_dict() for r in reports],
    }
    with open(output_path, "w") as f:
        json.dump(data, f, indent=2)
    print(f"[CONSISTENCY] Raw JSON report saved -> {output_path}")


# ----------------------------------------------------------------------
# Human-readable Word report
# ----------------------------------------------------------------------


def _needs_attention(report: ConsistencyReport) -> bool:
    """True if this section looks like the RLM struggled across runs."""
    j, m = report.judge_verdict, report.metrics
    if j and j.status_disagreement:
        return True
    if j and j.consistency_score is not None and j.consistency_score < 80:
        return True
    if m and not m.status_unanimous:
        return True
    if m and m.min_evidence_jaccard < 0.4:
        return True
    return False


def _score_color(score: int | None) -> RGBColor:
    if score is None:
        return RGBColor(0x80, 0x80, 0x80)
    if score >= 90:
        return RGBColor(0x1A, 0x7F, 0x37)
    if score >= 70:
        return RGBColor(0xB3, 0x7A, 0x00)
    return RGBColor(0xC0, 0x1C, 0x1C)


def _add_horizontal_rule(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_labeled_paragraph(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(f"{label} ")
    run.bold = True
    if text:
        for i, line in enumerate(text.split("\n")):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)
    else:
        italic = p.add_run("(empty)")
        italic.italic = True


def generate_consistency_docx(
    reports: list[ConsistencyReport],
    output_path: str,
    protocol_path: str,
    n_runs: int,
    model_name: str,
    seed: int | None = None,
    post_processing: dict[str, bool] | None = None,
) -> str:
    """Render a reviewer-friendly Word report: cover page, an executive-summary
    table flagging sections that look inconsistent, then one page per section
    with every run's status/confidence/filled text/evidence/notes stacked for
    visual side-by-side comparison, the LLM-judge verdict, and the quantitative
    metrics.
    """
    document = Document()

    # -- Cover page ------------------------------------------------------
    title = document.add_heading("RLM Extraction Consistency Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    post_processing = post_processing or {}
    meta_lines = [
        f"Protocol: {os.path.basename(protocol_path)}",
        f"Model: {model_name}",
        f"Runs compared: {n_runs}",
        f"Seed: {seed}",
        "Post-processing: "
        + ", ".join(f"{k}={'yes' if v else 'no'}" for k, v in post_processing.items()),
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Sections analyzed: {len(reports)}",
    ]
    for line in meta_lines:
        document.add_paragraph(line)

    document.add_paragraph(
        "\nHow to read this report: each section below shows the FINAL extraction "
        "text (after harmonization, plain-language review, and remediation) from "
        f"every one of the {n_runs} independent runs, stacked so you can compare them "
        "by eye. A quantitative agreement score and an LLM-as-judge verdict follow "
        "each set of runs. Sections flagged REVIEW in the summary table below "
        "are the ones most likely to have inconsistent extraction across runs."
    )

    # -- Executive summary -------------------------------------------------
    document.add_page_break()
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "One row per section. REVIEW means the runs disagreed on status, had low "
        "evidence overlap, or the LLM judge scored consistency below 80/100 — "
        "start there."
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Section", "Heading", "Status Agreement", "LLM Score", "Flag"], strict=True):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        _shade_cell(cell, "D9E2F3")

    for r in reports:
        row = table.add_row().cells
        row[0].text = r.section_id
        row[1].text = r.heading[:60]
        if r.metrics:
            row[2].text = (
                f"{r.metrics.status_agreement_rate:.0%} "
                f"({'unanimous' if r.metrics.status_unanimous else r.metrics.status_majority})"
            )
        score_run = row[3].paragraphs[0].add_run(
            f"{r.judge_verdict.consistency_score}/100" if r.judge_verdict else "N/A"
        )
        score_run.bold = True
        score_run.font.color.rgb = _score_color(
            r.judge_verdict.consistency_score if r.judge_verdict else None
        )
        if _needs_attention(r):
            flag_run = row[4].paragraphs[0].add_run("REVIEW")
            flag_run.bold = True
            flag_run.font.color.rgb = RGBColor(0xC0, 0x1C, 0x1C)
            _shade_cell(row[4], "FCE4E4")

    # -- Per-section detail --------------------------------------------------
    for r in reports:
        document.add_page_break()
        document.add_heading(f"[{r.section_id}] {r.heading}", level=1)

        j, m = r.judge_verdict, r.metrics
        verdict_p = document.add_paragraph()
        verdict_run = verdict_p.add_run(
            f"Consistency: {j.consistency_score}/100 ({j.consistency_label})"
            if j
            else "Consistency: N/A"
        )
        verdict_run.bold = True
        verdict_run.font.size = Pt(13)
        verdict_run.font.color.rgb = _score_color(j.consistency_score if j else None)
        if m:
            document.add_paragraph(
                f"Status agreement: {m.status_agreement_rate:.0%} "
                f"({'unanimous: ' + m.status_majority if m.status_unanimous else 'majority: ' + m.status_majority}) "
                f"  |  Evidence overlap (avg/min): {m.avg_evidence_jaccard:.2f} / {m.min_evidence_jaccard:.2f}"
            )

        # Quick-glance table
        glance = document.add_table(rows=1, cols=4)
        glance.style = "Light List Accent 1"
        for cell, text in zip(glance.rows[0].cells, ["Run", "Status", "Confidence", "Evidence #"], strict=True):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
        for run_result in r.runs:
            ext = run_result.extraction
            row = glance.add_row().cells
            row[0].text = f"Run {run_result.run_index}"
            row[1].text = ext.status
            row[2].text = ext.confidence
            row[3].text = str(len(ext.evidence))

        # Stacked per-run detail
        for run_result in r.runs:
            ext = run_result.extraction
            document.add_heading(f"Run {run_result.run_index}", level=2)
            document.add_paragraph(f"Status: {ext.status}    Confidence: {ext.confidence}")
            _add_labeled_paragraph(document, "Filled template:", ext.filled_template or ext.answer)
            if ext.evidence:
                p = document.add_paragraph()
                p.add_run("Evidence quotes:").bold = True
                for e in ext.evidence:
                    bullet = document.add_paragraph(style="List Bullet")
                    bullet.add_run(f'"{e.quote}" ').italic = True
                    bullet.add_run(f"(p. {e.page})")
            if ext.notes:
                _add_labeled_paragraph(document, "Notes:", ext.notes)
            _add_horizontal_rule(document)

        # LLM judge verdict
        if j:
            document.add_heading("LLM Judge Verdict", level=2)
            document.add_paragraph(f"Status disagreement flagged: {j.status_disagreement}")
            if j.facts_missing_in_some_runs:
                document.add_paragraph("Facts present in some runs but missing in others:").runs[
                    0
                ].bold = True
                for f in j.facts_missing_in_some_runs:
                    document.add_paragraph(
                        f"{f.get('fact')} — present in runs {f.get('present_in_runs')}, "
                        f"missing in runs {f.get('missing_in_runs')}",
                        style="List Bullet",
                    )
            if j.contradictions:
                document.add_paragraph("Contradictions between runs:").runs[0].bold = True
                for c in j.contradictions:
                    document.add_paragraph(
                        f"{c.get('description')} — runs {c.get('runs_involved')}",
                        style="List Bullet",
                    )
            _add_labeled_paragraph(document, "Judge summary:", j.summary)

        # Quantitative metrics
        if m:
            document.add_heading("Quantitative Metrics", level=2)
            metrics_lines = [
                f"Statuses: {m.statuses}    Confidences: {m.confidences}    "
                f"Evidence counts: {m.evidence_counts}",
                f"Exact text matches: {m.exact_match_pairs}/{m.total_pairs} run pairs byte-identical",
                f"Avg character similarity: {m.avg_char_similarity:.2f} (min {m.min_char_similarity:.2f})",
                f"Avg word overlap (Jaccard): {m.avg_word_jaccard:.2f} (min {m.min_word_jaccard:.2f})",
                f"Avg evidence-quote overlap (Jaccard): {m.avg_evidence_jaccard:.2f} "
                f"(min {m.min_evidence_jaccard:.2f})",
            ]
            for line in metrics_lines:
                document.add_paragraph(line, style="List Bullet")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    document.save(output_path)
    print(f"[CONSISTENCY] Word report saved -> {output_path}")
    return output_path
