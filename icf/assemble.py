"""
Output generation: marked-up ICF DOCX and JSON extraction report.

Produces two artefacts:
  1. marked_up_icf.docx       - A Word document following the template structure,
                                annotated with status, evidence, validation, and
                                review flags for traceability.
  2. extraction_report.json   - Full structured data for programmatic use.
"""

import json
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from icf.clean_icf import (
    MARKER_PLEASE_COMPLETE,
    _add_ai_disclaimer,
    _add_highlighted_placeholder,
    _resolve_section_placeholder_label,
    _section_suggested_text,
    _write_us_summary_page_opening,
)
from icf.runtime_injections import SIGNATURE_CONSENT_SECTION_ID
from icf.types import (
    ExtractionResult,
    RemediationResult,
    ReviewResult,
    TemplateVariable,
    ValidationResult,
)

# ------------------------------------------------------------------
# Colour constants
# ------------------------------------------------------------------
_GREY = RGBColor(128, 128, 128)
_RED = RGBColor(200, 30, 30)
_ORANGE = RGBColor(200, 130, 0)
_GREEN = RGBColor(30, 130, 30)


# ------------------------------------------------------------------
# 1. Marked-up ICF DOCX
# ------------------------------------------------------------------


def generate_marked_up_docx(
    extractions: list[ExtractionResult],
    validations: list[ValidationResult],
    variables: list[TemplateVariable],
    output_path: str,
    review_result: ReviewResult | None = None,
    us_funded: bool = False,
) -> str:
    """Create a marked-up DOCX with all sections, filled content, and markers."""
    doc = Document()

    # Title page
    title = doc.add_heading("MARKED-UP - Informed Consent Form", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph(
        "This is an auto-generated, marked-up draft. Sections highlighted in yellow "
        "require human review and completion. Evidence citations are included below "
        "each section for reference."
    )
    _style_run(note.runs[0], size=9, italic=True, colour=_GREY)

    doc.add_page_break()

    # Build quick lookup
    ext_map: dict[str, ExtractionResult] = {e.section_id: e for e in extractions}
    val_map: dict[str, ValidationResult] = {v.section_id: v for v in validations}

    if us_funded:
        _write_draft_us_summary(doc, variables, ext_map, val_map, review_result)
        doc.add_page_break()

    for var in variables:
        if var.section_id.startswith("1.") or var.section_id == SIGNATURE_CONSENT_SECTION_ID:
            continue
        ext = ext_map.get(var.section_id)
        val = val_map.get(var.section_id)

        # Omit optional sections that were not found / not applicable.
        # Required sections always appear so the human reviewer knows to fill them.
        if ext is not None and ext.status in ("NOT_FOUND", "SKIPPED") and not var.required:
            continue

        _write_draft_section(doc, var, ext, val, review_result)

    # Review appendix
    if review_result and (review_result.flags or review_result.cross_section_notes):
        doc.add_page_break()
        doc.add_heading("APPENDIX: Plain Language Review Flags", level=1)
        if review_result.cross_section_notes:
            note_p = doc.add_paragraph()
            note_r = note_p.add_run(f"Cross-section notes: {review_result.cross_section_notes}")
            _style_run(note_r, size=9, italic=True, colour=_GREY)
        if review_result.flags:
            high_flags = [f for f in review_result.flags if f.severity == "HIGH"]
            med_flags = [f for f in review_result.flags if f.severity == "MEDIUM"]
            low_flags = [f for f in review_result.flags if f.severity == "LOW"]
            for severity_group, colour in (
                (high_flags, _RED),
                (med_flags, _ORANGE),
                (low_flags, _GREY),
            ):
                for flag in severity_group:
                    bp = doc.add_paragraph(style="List Bullet")
                    br = bp.add_run(
                        f"[{flag.section_id}] {flag.issue_type} ({flag.severity}): "
                        f'"{flag.flagged_text[:120]}" — {flag.suggestion}'
                    )
                    _style_run(br, size=9, colour=colour)
                    if flag.suggested_fix:
                        sfp = doc.add_paragraph()
                        sfr = sfp.add_run(f"    Suggested fix: {flag.suggested_fix}")
                        _style_run(sfr, size=9, colour=colour)

    _add_ai_disclaimer(doc)

    doc.save(output_path)
    return output_path


# ------------------------------------------------------------------
# 2. JSON report
# ------------------------------------------------------------------


def generate_report_json(
    extractions: list[ExtractionResult],
    validations: list[ValidationResult],
    summary: dict[str, Any],
    output_path: str,
    review_result: ReviewResult | None = None,
    remediation_result: RemediationResult | None = None,
) -> str:
    """Write the full extraction report as JSON."""
    report = {
        "summary": summary,
        "extractions": [e.to_dict() for e in extractions],
        "validations": [v.to_dict() for v in validations],
        "review": review_result.to_dict() if review_result else None,
        "remediation": remediation_result.to_dict() if remediation_result else None,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)
    return output_path


# ------------------------------------------------------------------
# US-funded summary (sections 1.x)
# ------------------------------------------------------------------


def _write_draft_us_summary(
    doc: Document,
    variables: list[TemplateVariable],
    ext_map: dict[str, ExtractionResult],
    val_map: dict[str, ValidationResult],
    review_result: ReviewResult | None,
) -> None:
    """Draft annotated Summary of ICF block — same layout as clean/validation docs."""
    _write_us_summary_page_opening(doc, ext_map)

    summary_vars = [v for v in variables if v.section_id.startswith("1.")]
    for var in summary_vars:
        _write_draft_section(
            doc,
            var,
            ext_map.get(var.section_id),
            val_map.get(var.section_id),
            review_result,
            include_heading=False,
        )


def _write_draft_section(
    doc: Document,
    var: TemplateVariable,
    ext: ExtractionResult | None,
    val: ValidationResult | None,
    review_result: ReviewResult | None,
    *,
    include_heading: bool = True,
) -> None:
    """Render one section in the draft ICF (shared by main body and US summary)."""
    if ext is not None and ext.status in ("NOT_FOUND", "SKIPPED") and not var.required:
        return

    if include_heading:
        level = 2 if var.sub_section else 1
        heading_text = var.heading
        if var.sub_section:
            heading_text += f" - {var.sub_section}"
        doc.add_heading(heading_text, level=level)
    elif var.section_id != "1.1" and var.sub_section:
        p = doc.add_paragraph()
        r = p.add_run(var.sub_section)
        _style_run(r, size=11, bold=True)

    if ext is None:
        _add_status_line(doc, "NOT PROCESSED", _GREY)
        return

    colour = {
        "FOUND": _GREEN,
        "PARTIAL": _ORANGE,
        "STANDARD_TEXT": _GREEN,
        "NOT_FOUND": _RED,
        "SKIPPED": _GREY,
        "ERROR": _RED,
    }.get(ext.status, _GREY)
    badge = f"Status: {ext.status}"
    if ext.confidence and ext.confidence != "N/A":
        badge += f"  |  Confidence: {ext.confidence}"
    if ext.error:
        badge += f"  |  Error: {ext.error}"
    _add_status_line(doc, badge, colour)

    if ext.status in ("FOUND", "PARTIAL", "STANDARD_TEXT"):
        text = ext.filled_template or ext.answer
        if text:
            _add_content(doc, text)
        if ext.status == "PARTIAL" and ext.notes:
            p = doc.add_paragraph()
            r = p.add_run(f"[PARTIAL] {ext.notes}")
            _style_run(r, size=9, colour=_ORANGE, italic=True)

    elif ext.status in ("NOT_FOUND", "SKIPPED"):
        label = _resolve_section_placeholder_label(var, optional=not var.required)
        _add_highlighted_placeholder(doc, label)
        suggested = _section_suggested_text(var)
        if suggested:
            sg = doc.add_paragraph()
            sr = sg.add_run("Suggested text: " + suggested[:800])
            _style_run(sr, size=10, colour=_GREY, italic=True)

    elif ext.status == "ERROR":
        p = doc.add_paragraph()
        r = p.add_run(f"[EXTRACTION ERROR] {ext.error}")
        _style_run(r, size=10, colour=_RED, bold=True)

    if ext.evidence:
        ep = doc.add_paragraph()
        er = ep.add_run("Evidence:")
        _style_run(er, size=8, italic=True, colour=_GREY)
        for ev in ext.evidence:
            bp = doc.add_paragraph(style="List Bullet")
            short_quote = ev.quote[:250].replace("\n", " ")
            br = bp.add_run(f'Page {ev.page}: "{short_quote}"')
            _style_run(br, size=8, italic=True, colour=_GREY)

    if val and val.issues:
        for issue in val.issues:
            ip = doc.add_paragraph()
            ir = ip.add_run(f"[VALIDATION] {issue}")
            _style_run(ir, size=8, colour=_ORANGE)

    if review_result:
        section_flags = [f for f in review_result.flags if f.section_id == var.section_id]
        for flag in section_flags:
            flag_colour = {"HIGH": _RED, "MEDIUM": _ORANGE, "LOW": _GREY}.get(
                flag.severity, _GREY
            )
            fp = doc.add_paragraph()
            fr = fp.add_run(
                f"[REVIEW | {flag.issue_type} | {flag.severity}] "
                f'"{flag.flagged_text[:120]}" — {flag.suggestion}'
            )
            _style_run(fr, size=8, colour=flag_colour, italic=True)
            if flag.suggested_fix:
                sfp = doc.add_paragraph()
                sfr = sfp.add_run(f"    Suggested fix: {flag.suggested_fix}")
                _style_run(sfr, size=8, colour=flag_colour)


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------


# ------------------------------------------------------------------
# Markdown table helpers
# ------------------------------------------------------------------

_TABLE_ROW_RE = re.compile(r"^\s*\|")
_TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-|:]+\|\s*$")


def _split_content(text: str) -> list[tuple[str, str]]:
    """Split text into ('text', ...) and ('table', ...) segments.

    Consecutive lines beginning with '|' are treated as a Markdown table block.
    All other lines form plain-text segments.
    """
    segments: list[tuple[str, str]] = []
    current_kind: str | None = None
    current_lines: list[str] = []

    for line in text.split("\n"):
        kind = "table" if _TABLE_ROW_RE.match(line) else "text"
        if kind != current_kind:
            if current_lines:
                segments.append((current_kind, "\n".join(current_lines)))  # type: ignore[arg-type]
            current_kind = kind
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_lines and current_kind is not None:
        segments.append((current_kind, "\n".join(current_lines)))

    return segments


def _parse_markdown_table(table_text: str) -> list[list[str]]:
    """Return a list of rows (each a list of cell strings) from a Markdown table.

    Separator lines (|---|---| etc.) are skipped.  The first data row is the header.
    """
    rows: list[list[str]] = []
    for line in table_text.strip().split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            continue
        if _TABLE_SEP_RE.match(line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_docx_table(doc: Document, rows: list[list[str]], colour: RGBColor | None = None) -> None:
    """Insert a bordered Word table from parsed Markdown rows.

    The first row is rendered as a bold header.
    """
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for ri, row_data in enumerate(rows):
        is_header = ri == 0
        for ci in range(n_cols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = is_header
            if colour is not None:
                run.font.color.rgb = colour
            # Shade header row light grey
            if is_header:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    from docx.oxml import OxmlElement

                    shd = OxmlElement("w:shd")
                    tc_pr.append(shd)
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")

    # Spacing after table
    doc.add_paragraph()


def _add_content_runs(para, text: str) -> None:
    """Render body text, highlighting [PLEASE COMPLETE] in bold yellow."""
    if MARKER_PLEASE_COMPLETE not in text:
        r = para.add_run(text)
        _style_run(r, size=11)
        return
    parts = text.split(MARKER_PLEASE_COMPLETE)
    for i, part in enumerate(parts):
        if part:
            r = para.add_run(part)
            _style_run(r, size=11)
        if i < len(parts) - 1:
            r = para.add_run(MARKER_PLEASE_COMPLETE)
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.bold = True
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_content(doc: Document, text: str) -> None:
    """Render text that may contain embedded Markdown tables into the document."""
    for kind, segment in _split_content(text):
        if kind == "table":
            rows = _parse_markdown_table(segment)
            if rows:
                _add_docx_table(doc, rows)
                continue
        for line in segment.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            p = doc.add_paragraph()
            _add_content_runs(p, stripped)


def _add_status_line(doc: Document, text: str, colour: RGBColor) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _style_run(r, size=8, colour=colour)


def _style_run(
    run,
    size: int | None = None,
    bold: bool = False,
    italic: bool = False,
    colour: RGBColor | None = None,
) -> None:
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour is not None:
        run.font.color.rgb = colour
